# -*- coding: utf-8 -*-
"""轻量 Markdown → DOCX 转换器：专用于 report/工作总结.md。
支持：标题/表格/列表/粗体/行内码/引用块/图片(![](相对路径)，相对 md 所在目录解析)。"""
import os
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = r"D:\cursorproject\intern\report\阶段三\阶段三_04-15转载机错配事件深度核查.md"
DST = r"D:\cursorproject\intern\report\阶段三\阶段三_04-15转载机错配事件深度核查.docx"
IMG_WIDTH_CM = 14.0  # A4 正文可用宽度约 16cm

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
CODE_RE = re.compile(r"`([^`]+)`")
IMG_RE = re.compile(r"^\s*!\[([^\]]*)\]\(([^)\s]+)\)\s*$")


def add_runs(par, text):
    """处理 **粗体** 与 `行内码`。"""
    parts = BOLD_RE.split(text)
    for i, seg in enumerate(parts):
        if seg == "":
            continue
        if i % 2 == 1:  # 粗体段
            r = par.add_run(seg)
            r.bold = True
        else:
            subs = CODE_RE.split(seg)
            for j, sub in enumerate(subs):
                if sub == "":
                    continue
                r = par.add_run(sub)
                if j % 2 == 1:
                    r.font.name = "Consolas"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")


def set_cn(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_picture_line(doc, path):
    """嵌入 ![](path)，path 相对 md 所在目录解析，居中显示。"""
    full = os.path.join(os.path.dirname(os.path.abspath(SRC)), path)
    if not os.path.exists(full):
        print(f"[warn] 图片不存在，跳过: {full}")
        return
    doc.add_picture(full, width=Cm(IMG_WIDTH_CM))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    lines = [l.rstrip("\n") for l in open(SRC, encoding="utf-8")]
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("---"):
            i += 1
            continue
        if not line.strip():
            i += 1
            continue

        # 表格：收集连续表格行
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:\-|]+\|?$", lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            ncols = len(header)
            tbl = doc.add_table(rows=1, cols=ncols)
            tbl.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                add_runs(tbl.rows[0].cells[j].paragraphs[0], h)
                for r in tbl.rows[0].cells[j].paragraphs[0].runs:
                    set_cn(r, size=10, bold=True)
            for row in rows:
                cells = tbl.add_row().cells
                for j in range(ncols):
                    if j < len(row):
                        add_runs(cells[j].paragraphs[0], row[j])
                    for r in cells[j].paragraphs[0].runs:
                        set_cn(r, size=10)
            doc.add_paragraph()
            continue

        # 图片
        m = IMG_RE.match(line)
        if m:
            add_picture_line(doc, m.group(2))
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            p = doc.add_paragraph()
            r = p.add_run(m.group(2))
            sizes = {1: 18, 2: 15, 3: 12.5, 4: 11.5, 5: 11, 6: 10.5}
            set_cn(r, size=sizes.get(level, 11), bold=True, color=(31, 78, 121) if level <= 3 else None)
            i += 1
            continue

        # 引用块
        if line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_runs(p, line.lstrip(">").strip())
            for r in p.runs:
                set_cn(r, size=10, color=(89, 89, 89))
            i += 1
            continue

        # 列表项
        m = re.match(r"^\s*(?:[-*]|\d+\.)\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet" if not re.match(r"^\s*\d", line) else "List Number")
            add_runs(p, m.group(1))
            for r in p.runs:
                set_cn(r, size=10.5)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        add_runs(p, line)
        for r in p.runs:
            set_cn(r, size=10.5)
        i += 1

    doc.save(DST)
    print(f"OK -> {DST}")


if __name__ == "__main__":
    sys.exit(main())
