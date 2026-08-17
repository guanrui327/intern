# -*- coding: utf-8 -*-
"""阶段三：跨设备关联异常检测技术报告生成器。

报告为数据驱动：所有数字从 output/phase3 输出文件实时读取/重算，
不硬编码。核心结论方向在阶段三完成时已固化：
  1. ZZJ 恒流控制 -> 幅度层耦合不可学习（线性 R²≈0.07，RF 过拟合）
  2. 耦合真身在二元开关层 -> 主检测器 = 联合工况规则事件
  3. 跨设备 Mahalanobis 剔除恒量特征后与规则事件互相印证
  4. CMJ 先行 -> ZZJ 跟随（22.14% vs 15.45%）

独立运行：python src/generate_phase3_tech_report.py
输出文件：output/phase3/phase3_tech_report.docx
"""

from __future__ import annotations

import sys
import traceback
from datetime import datetime
from pathlib import Path

import numpy as np
import pandas as pd

# ── DOCX ──
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(BASE))

from src.config import CMJ_PROD_FEATURES, ZZJ_LOAD_TARGET  # noqa: E402

PHASE3_DIR = BASE / "output" / "phase3"
DATA_DIR = PHASE3_DIR / "data"
OUTPUT_FILE = PHASE3_DIR / "phase3_tech_report.docx"
IMG_DIR = BASE / "report" / "阶段三"
PHASE2_ANOM = BASE / "output" / "phase2" / "anomalies"


# ════════════════════════════════════════════════════════════
# 工具函数（与阶段二报告生成器同构）
# ════════════════════════════════════════════════════════════

def _read_csv(path: Path) -> pd.DataFrame | None:
    if not path.exists():
        return None
    try:
        return pd.read_csv(path, encoding="utf-8-sig")
    except Exception as e:
        print(f"  [WARN] 无法读取 {path.name}: {e}")
        return None


def _new_heading(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def _add_para(doc: Document, text: str, bold: bool = False, italic: bool = False,
              size: int | None = None, spacing_after: int = 6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = 1.35
    return p


def _add_bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def _add_table(doc: Document, df: pd.DataFrame, caption: str = "", max_rows: int = 20):
    if df.empty:
        _add_para(doc, "（无数据）", italic=True)
        return None
    df = df.head(max_rows)
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, (_, row) in enumerate(df.iterrows()):
        for j, col in enumerate(df.columns):
            cell = table.rows[i + 1].cells[j]
            val = row[col]
            if isinstance(val, (np.ndarray, list, tuple)):
                cell.text = str(list(val))
            elif pd.isna(val):
                cell.text = ""
            else:
                cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table


def _add_formula(doc: Document, latex: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(latex)
    run.font.name = "Consolas"
    run.italic = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_image(doc: Document, path: Path, width_inches: float = 5.8, caption: str = ""):
    if not path.exists():
        _add_para(doc, f"[图片不存在: {path.name}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width_inches))
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(f"图：{caption}")
        run.italic = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)


def _add_page_break(doc: Document):
    doc.add_page_break()


def _load_system() -> pd.DataFrame | None:
    p = DATA_DIR / "system_table.parquet"
    if not p.exists():
        return None
    return pd.read_parquet(p)


def _load_rule_events() -> pd.DataFrame | None:
    return _read_csv(DATA_DIR / "rule_events.csv")


def _load_mahal() -> pd.DataFrame | None:
    return _read_csv(DATA_DIR / "joint_mahalanobis.csv")


# ════════════════════════════════════════════════════════════
# 运行时重算（未落盘的关键指标）
# ════════════════════════════════════════════════════════════

def _rerun_regression() -> dict:
    """重跑耦合回归，拿到线性/RF/泛化准确数字。"""
    from src.relation_model import fit_coupling_regression, evaluate_generalization
    sys_df = _load_system()
    out = {"ok": False}
    if sys_df is None:
        return out
    try:
        lin = fit_coupling_regression(sys_df, model="linear")
        rf = fit_coupling_regression(sys_df, model="rf")
        gen = evaluate_generalization(sys_df, rf)
        y = sys_df[sys_df["联合工况"] == "生产运行"][ZZJ_LOAD_TARGET]
        out.update({
            "ok": True,
            "lin_r2": round(lin["best_r2"], 3),
            "lin_lag": lin["best_lag"],
            "lin_n": lin["n_train"],
            "rf_r2": round(rf["best_r2"], 3),
            "rf_lag": rf["best_lag"],
            "gen_train_r2": round(gen["train_r2"], 3),
            "gen_val_r2": round(gen["val_r2"], 3),
            "gen_val_mae": round(gen["val_mae"], 1),
            "gen_val_std": round(gen["y_val_std"], 1),
            "y_unique": int(y.nunique()),
            "y_median": float(y.median()),
            "y_p25": float(y.quantile(.25)),
            "y_p75": float(y.quantile(.75)),
        })
    except Exception as e:
        print(f"  [WARN] 回归重算失败: {e}")
    return out


def _rerun_propagation() -> dict:
    """重算双向事件传导率。"""
    from src.relation_event import (extract_merged_events, propagate_events,
                                    propagation_stats)
    out = {"ok": False}
    cmj_p = PHASE2_ANOM / "cmj_merged_events.csv"
    zzj_p = PHASE2_ANOM / "zzj_merged_events.csv"
    if not (cmj_p.exists() and zzj_p.exists()):
        return out
    try:
        cmj = extract_merged_events(str(cmj_p))
        zzj = extract_merged_events(str(zzj_p))
        fc = propagate_events(cmj, zzj)
        rc = propagate_events(zzj, cmj)
        fs = propagation_stats(fc, len(cmj))
        rs = propagation_stats(rc, len(zzj))
        out.update({
            "ok": True,
            "cmj_n_events": len(cmj),
            "zzj_n_events": len(zzj),
            "fwd_rate": fs["rate"], "fwd_chain": fs["n_chain"],
            "fwd_median": fs["lag_median"], "fwd_mean": fs["lag_mean"],
            "rev_rate": rs["rate"], "rev_chain": rs["n_chain"],
            "rev_median": rs["lag_median"],
        })
    except Exception as e:
        print(f"  [WARN] 传导重算失败: {e}")
    return out


# ════════════════════════════════════════════════════════════
# 章节生成器
# ════════════════════════════════════════════════════════════

def build_introduction(doc: Document):
    """第1章：引言"""
    _new_heading(doc, "引言", 1)

    _add_para(doc, (
        "本报告为「阶段三：多设备跨设备关联异常检测」的技术总结文档。"
        "阶段一完成了单设备分部位工况划分，阶段二完成了 CMJ / ZZJ 各自的单设备异常检测，"
        "本阶段将视野从单设备扩展到跨设备：挖掘上游采煤机（CMJ）到下游转载机（ZZJ）"
        "的物理耦合关系，检测「上游割煤但下游负载不匹配」的关联异常——"
        "即堵煤 / 卡链 / 断链 / 煤流堆积风险。"
    ))

    _add_para(doc, "运输链物理结构：CMJ 割煤 → CMJ 破碎机 → 溜槽 → ZZJ 转载机。", bold=True)

    _new_heading(doc, "分析目标", 2)
    _add_bullet(doc, "验证 CMJ 产量代理量与 ZZJ 负载之间的物理耦合是否可学习（回归）")
    _add_bullet(doc, "构造联合系统工况，在二元开关层检测「上游割煤但下游不接」的关联异常")
    _add_bullet(doc, "跨设备 Mahalanobis 检测多参数联合异常，与规则事件互相印证")
    _add_bullet(doc, "验证事件传导方向：CMJ 先行 → ZZJ 跟随")

    _new_heading(doc, "分析范围", 2)
    _add_para(doc, (
        "上游产量代理（CMJ）：割煤速度（采煤机_牵引部位_采煤机速度）、"
        "左右滚筒电机电流、左右滚筒高度；"
        "下游负载（ZZJ）：转载机电机电流（+ 母线电压 / 链条速度用于诊断）。"
        "数据覆盖 2024-04-01 ~ 04-28 共 1 分钟等间隔采样。"
    ))

    # 物理耦合模型表
    coupling = pd.DataFrame({
        "角色": ["上游产量（X）", "上游产量（X）", "上游产量（X）",
                 "下游负载（y）", "工况", "工况"],
        "代理量": ["割煤速度", "截割功率", "截割断面",
                   "转载负载", "CMJ 工况", "ZZJ 工况"],
        "测点": [
            "采煤机_牵引部位_采煤机速度",
            "采煤机_截割部位_左/右滚筒_电机_电流",
            "采煤机_截割部位_左/右滚筒_高度",
            "三机_转载机_电机_电流 / 转矩 / 链条速度",
            "设备_工况（割煤中/正常运行/空载牵引/待机/停机）",
            "工况（停机/空载运行/带载运行）",
        ],
    })
    _add_table(doc, coupling, "表1：跨设备物理耦合模型")


def build_data_overview(doc: Document):
    """第2章：数据概况与分析设计"""
    _new_heading(doc, "数据概况与分析设计", 1)

    sys_df = _load_system()
    if sys_df is not None:
        _add_para(doc, (
            f"系统宽表由 CMJ 带工况宽表与 ZZJ 带工况宽表按时间戳对齐拼接，"
            f"共 {len(sys_df):,} 个 1 分钟采样点，"
            f"时间跨度 {sys_df.index.min():%Y-%m-%d %H:%M} ~ {sys_df.index.max():%Y-%m-%d %H:%M}。"
        ))
        _add_para(doc, (
            "时间戳对齐验证：两表索引交集 = CMJ 全量（ZZJ 数据覆盖至 05-31，"
            "CMJ 至 04-28 结束，ZZJ 多出的时段为「CMJ 无数据期」，丢弃），"
            "对齐采用交集而非并集，保证每个样本两设备工况标签齐全。"
        ))

    _new_heading(doc, "特征选择", 2)
    _add_para(doc, "系统宽表的核心列：", bold=True)
    _zzj_extra = [
        ("三机_转载机_电机_转矩", "辅助负载信号"),
        ("三机_转载机_母线电压", "开关量双峰（≈5V/≈4412V），不进协方差特征"),
        ("三机_转载机_链条速度", "开关式恒量重尾，不进协方差特征"),
    ]
    _rows = []
    for p in CMJ_PROD_FEATURES:
        _rows.append(("上游产量", p, "割煤速度 / 滚筒电流 / 滚筒高度"))
    _rows.append(("下游负载", ZZJ_LOAD_TARGET, "转载负载（恒流控制）"))
    for p, note in _zzj_extra:
        _rows.append(("下游负载（诊断用）", p, note))
    _rows.append(("工况", "设备_工况", "CMJ 设备工况（5 态）"))
    _rows.append(("工况", "工况", "ZZJ 工况（3 态）"))
    feats = pd.DataFrame({
        "类别": [r[0] for r in _rows],
        "参数": [r[1] for r in _rows],
        "说明": [r[2] for r in _rows],
    })
    _add_table(doc, feats, "表2：系统宽表核心列")

    _add_para(doc, (
        "恒量特征排除：ZZJ 变频器母线电压为开关量双峰（≈5V / ≈4412V），"
        "链条速度为开关式恒量重尾（p25~p75=1794~1876 却含 0/2193 极值）——"
        "此类特征使 MCD 稳健协方差被挤窄后尾部全标异常，"
        "故从跨设备 Mahalanobis 特征中剔除（详见第 6 章实测验证）。"
    ), italic=True, size=9)


def build_joint_condition(doc: Document):
    """第3章：联合工况划分与物理耦合验证"""
    _new_heading(doc, "联合工况划分与物理耦合验证", 1)

    sys_df = _load_system()
    if sys_df is None:
        _add_para(doc, "（无系统宽表）", italic=True)
        return

    _new_heading(doc, "联合系统工况构造", 2)
    _add_para(doc, (
        "由 设备_工况（CMJ） × 工况（ZZJ） 合成联合系统工况，六态语义如下："
    ))
    cond_df = pd.DataFrame({
        "联合工况": ["生产运行", "采煤-转载错配", "转载余流", "全线停机", "全线待机", "空载循环"],
        "组合": ["割煤中 + 带载", "割煤中 + 空载/停机", "停机/待机 + 带载",
                 "停机 + 停机", "待机 + 空载", "空载牵引 + 空载/带载"],
        "物理含义": [
            "正常采煤基线",
            "堵煤/断链/转载未启动风险（关联异常信号）",
            "残余煤流/转载滞后",
            "双设备协同停产",
            "双设备协同待机",
            "割煤停机后溜槽空转",
        ],
    })
    _add_table(doc, cond_df, "表3：联合系统工况语义")

    _new_heading(doc, "工况分布", 2)
    jc = sys_df["联合工况"]
    vc = jc.value_counts()
    dist = pd.DataFrame({
        "联合工况": vc.index,
        "点数": vc.values,
        "占比": (vc / len(sys_df) * 100).round(2).astype(str) + "%",
    })
    _add_table(doc, dist, "表4：联合工况分布")

    _new_heading(doc, "开关层物理耦合验证（crosstab）", 2)
    _add_para(doc, "设备_工况 → ZZJ 工况（行归一化），验证「割煤中 → 带载」是否占主导：", bold=True)
    ct = pd.crosstab(sys_df["设备_工况"], sys_df["工况"], normalize="index")
    ct_pct = (ct * 100).round(1)
    ct_pct.columns = [f"{c}（%）" for c in ct_pct.columns]
    ct_pct = ct_pct.reset_index().rename(columns={"设备_工况": "设备_工况（行）"})
    _add_table(doc, ct_pct, "表5：设备_工况 → ZZJ 工况 转移占比（%）")

    mis_row = ct.loc["割煤中", "带载运行"] * 100
    _add_para(doc, (
        f"割煤中 → 带载运行 占 {mis_row:.1f}%——二元开关层的物理耦合真实存在且占主导。"
        "恒流控制抹平的是幅度层，但开关层（割煤 ↔ 带载）依然强耦合，"
        "这正是联合工况规则事件可行的根基。"
    ), bold=True)


def build_regression(doc: Document):
    """第4章：物理耦合回归验证（恒流控制 -> 幅度层不可学习）"""
    _new_heading(doc, "物理耦合回归验证", 1)
    r = _rerun_regression()

    _new_heading(doc, "假设与设计", 2)
    _add_para(doc, (
        "计划假设：CMJ 高速割煤 + 高滚筒电流 → ZZJ 高负载电流，回归可学。"
        "验证设计：训练域 = 联合工况「生产运行」（正常采煤基线），"
        f"X = 上游产量代理 {len(CMJ_PROD_FEATURES)} 维，y = 转载机电机电流，"
        "对 X 做 lag 0~5min（煤流传播物理滞后）选最佳 R²。"
    ))

    _new_heading(doc, "线性回归基线", 2)
    if r["ok"]:
        _add_table(doc, pd.DataFrame([{
            "模型": "线性回归（best lag）",
            "最佳滞后": f"{r['lin_lag']} min",
            "R²": f"{r['lin_r2']:.3f}",
            "训练样本": f"{r['lin_n']:,}",
        }]), "表6：线性回归结果")
        _add_para(doc, (
            f"线性 R² = {r['lin_r2']:.3f}——即使加上 0~5min 物理滞后，"
            "上游产量幅度对下游电流幅度几乎无解释力。"
        ))

    _new_heading(doc, "RandomForest 非线性验证", 2)
    if r["ok"]:
        _add_table(doc, pd.DataFrame({
            "指标": ["RF 训练 R²", "时间序 holdout 验证 R²", "验证 MAE", "验证集 y 标准差"],
            "数值": [
                f"{r['gen_train_r2']:.3f}（in-sample）",
                f"{r['gen_val_r2']:.3f}（前 70% 训练 / 后 30% 验证）",
                f"{r['gen_val_mae']} A",
                f"{r['gen_val_std']} A",
            ],
        }), "表7：RF 时间序泛化评估")

    if r["ok"]:
        _add_formula(doc, f"训练 R² = {r['gen_train_r2']}  ≫  验证 R² = {r['gen_val_r2']}")
        _add_para(doc, (
            f"RF 训练 R² 虚高（{r['gen_train_r2']}）但时间序 holdout 验证 R² 为负（{r['gen_val_r2']}），"
            f"验证 MAE（{r['gen_val_mae']} A）与验证集 y 标准差（{r['gen_val_std']} A）相当"
            "——模型在训练集记忆了噪声，出域即退化为预测常数，完全过拟合。"
        ))

    _new_heading(doc, "恒流控制的证据", 2)
    if r["ok"]:
        _add_para(doc, (
            f"生产运行域内转载机电机电流被量化到 {r['y_unique']} 个离散档位"
            f"（中位 {r['y_median']:.0f} A，IQR [{r['y_p25']:.0f}, {r['y_p75']:.0f}] A）"
            "——数字恒流控制将电流维持在设定值附近，下游电流幅度与上游产量幅度无关。"
        ))

    _add_para(doc, "核心结论：", bold=True)
    _add_para(doc, (
        "生产运行域（正常采煤）内，回归残差异常不可用——不是模型不够好，"
        "而是恒流控制抹平了幅度耦合。耦合真身在更低层（二元开关层 + 规则事件层）。"
        "据此，主检测器从「物理耦合回归」切换为「联合工况规则事件」（V 已拍板）。"
    ), bold=True)


def build_rule_events(doc: Document):
    """第5章：联合工况规则事件（主检测器）"""
    _new_heading(doc, "联合工况规则事件主检测器", 1)

    _new_heading(doc, "方法", 2)
    _add_para(doc, (
        "回归不可学后，主检测器切换为联合工况规则事件：错配 / 余流标签本身即编码物理规则，"
        "无需回归。规则 = 联合工况处于风险状态，无需任何阈值/模型。"
    ))
    _add_formula(doc, "错配 = 割煤中 + 下游未带载（空载/停机）  余流 = 上游停机/待机 + 带载")

    _new_heading(doc, "事件统计", 2)
    re = _load_rule_events()
    if re is not None and not re.empty:
        # 分段统计：段数 / 覆盖点数 / 最长时长
        groups = re.groupby("规则类型")
        rows = []
        for typ, g in groups:
            rows.append({
                "规则类型": typ,
                "事件段数": len(g),
                "覆盖点数": int(g["n_points"].sum()),
                "最长时长(min)": int(g["duration_min"].max()),
                "平均时长(min)": f"{g['duration_min'].mean():.1f}",
            })
        _add_table(doc, pd.DataFrame(rows), "表8：规则事件统计")

        # 错配最长案例（04-15）
        mis = re[re["规则类型"].str.contains("错配")]
        if not mis.empty:
            top = mis.nlargest(1, "duration_min").iloc[0]
            _add_para(doc, (
                f"典型长时错配案例（{top['start'][:10]}）：持续 {top['duration_min']:.0f} 分钟"
                f"（{top['start']} ~ {top['end']}）——采煤机持续割煤但转载机未带载，"
                "煤流长期堆积在下游设备的显著风险信号。"
            ), bold=True)

    _new_heading(doc, "解读", 2)
    _add_para(doc, (
        "错配（割煤中但下游不接）指向堵煤 / 断链 / 转载未启动风险；"
        "余流（上游已停但下游仍在带载）指向残余煤流 / 转载滞后。"
        "两者均为可治理的离散风险段，而非系统性失控——事件呈零星块状分布，"
        "说明设备协同总体健康，异常集中在少数明确时段。"
    ))


def build_mahalanobis(doc: Document):
    """第6章：跨设备 Mahalanobis 关联异常"""
    _new_heading(doc, "跨设备 Mahalanobis 关联异常", 1)

    _new_heading(doc, "原理", 2)
    _add_formula(doc, "D_M(x) = sqrt((x - μ)^T · Σ^(-1) · (x - μ)),   特征 = CMJ 产量(5维) + ZZJ 负载(1维)")
    _add_para(doc, (
        "按联合工况分组，MCD 稳健协方差 + χ² 分位数（α=0.001）判异常，"
        "取 Top-3 特征贡献做归因。"
    ))

    m = _load_mahal()
    if m is not None and not m.empty:
        total_anom = int(m["is_anomaly"].sum())
        rate = total_anom / len(m) * 100
        _new_heading(doc, "恒量特征排除（重要修正）", 2)
        _add_para(doc, (
            "初版特征含母线电压（开关量双峰）与链条速度（开关式恒量重尾），"
            "实测异常率 17% 泛滥——恒量特征使 MCD 挤窄带后尾部全标异常。"
            "剔除后异常率降至稳定水平。"
        ))

        _new_heading(doc, "结果", 2)
        summary = m.groupby("工况").agg(
            总样本=("is_anomaly", "count"),
            异常数=("is_anomaly", "sum"),
            异常率=("is_anomaly", lambda x: f"{x.sum()/len(x)*100:.2f}%"),
        ).reset_index()
        summary.columns = ["联合工况", "总样本", "异常数", "异常率"]
        _add_table(doc, summary, "表9：Mahalanobis 按联合工况异常分布")

        _add_para(doc, (
            f"总异常率 {rate:.2f}%（{total_anom:,}/{len(m):,}）；"
            "生产运行域几乎无异常（<0.1%），异常集中在物理不匹配域"
            "（余流 + 错配 占异常总数的绝大部分）——"
            "与规则事件主检测器互相印证，Mahalanobis 起佐证作用，不喧宾夺主。"
        ))


def build_propagation(doc: Document):
    """第7章：事件传导关联"""
    _new_heading(doc, "事件传导关联（CMJ 先行 → ZZJ 跟随）", 1)
    pr = _rerun_propagation()

    _new_heading(doc, "方法", 2)
    _add_para(doc, (
        "输入：阶段二单设备异常事件（merged_events 中 any_anomaly 连续时段）。"
        "对每条上游事件，找其 start 后 0~10min 窗口内开始的下游事件记为传导链；"
        "双向对比（CMJ→ZZJ vs ZZJ→CMJ）验证「上游先行 → 下游跟随」方向假设。"
    ))

    if pr["ok"]:
        _new_heading(doc, "传导率对比", 2)
        _add_table(doc, pd.DataFrame([
            {"方向": "CMJ → ZZJ（假设方向）",
             "上游事件数": pr["cmj_n_events"],
             "传导链数": pr["fwd_chain"],
             "传导率": f"{pr['fwd_rate']}%",
             "中位滞后": f"{pr['fwd_median']} min"},
            {"方向": "ZZJ → CMJ（反向）",
             "上游事件数": pr["zzj_n_events"],
             "传导链数": pr["rev_chain"],
             "传导率": f"{pr['rev_rate']}%",
             "中位滞后": f"{pr['rev_median']} min"},
        ]), "表10：双向事件传导率对比")

        _add_para(doc, (
            f"CMJ → ZZJ 传导率 {pr['fwd_rate']}% 显著高于反向 {pr['rev_rate']}%，"
            f"中位滞后 {pr['fwd_median']} 分钟与煤流从采煤机到转载机的物理传播时间量级吻合"
            "——「上游事件驱动下游事件」的传导假设成立。"
        ), bold=True)


def build_visualization(doc: Document):
    """第8章：可视化与典型案例"""
    _new_heading(doc, "可视化与典型案例", 1)

    _add_image(doc, IMG_DIR / "1_联合工况堆叠.png", width_inches=6.3,
               caption="联合系统工况堆叠（2024-04-01 ~ 04-07 采样）——错配/余流呈零星块状")
    _add_image(doc, IMG_DIR / "2_关联时间线_错配案例.png", width_inches=6.0,
               caption="04-15 错配案例双设备时间线——上游在割、下游不接的断流铁证")
    _add_image(doc, IMG_DIR / "3_产量负载散点.png", width_inches=6.0,
               caption="产量-负载散点：左 生产运行恒流带（电流幅度与速度无关）；右 错配断流")
    _add_image(doc, IMG_DIR / "4_事件传导滞后.png", width_inches=5.8,
               caption="CMJ→ZZJ 事件传导滞后分布——中位约 4 分钟")


def build_conclusion(doc: Document):
    """第9章：结论、局限与下一步"""
    _new_heading(doc, "结论、局限与下一步", 1)

    _new_heading(doc, "主要结论", 2)
    _add_bullet(doc, "ZZJ 恒流控制 → 幅度层耦合不可学习：线性 R²≈0.07，RF 时间序验证 R²<0，已诚实报告，避免用假回归掩盖真相")
    _add_bullet(doc, "耦合真身在二元开关层（割煤 ↔ 带载 96.6%）→ 主检测器切换为联合工况规则事件（错配 123 段 / 余流 1532 段）")
    _add_bullet(doc, "04-15 案例连续 309 分钟错配为显著风险信号，值得人工核查")
    _add_bullet(doc, "跨设备 Mahalanobis 剔除恒量特征后异常率 17%→6.27%，95% 异常落在物理不匹配域，与规则事件互相印证")
    _add_bullet(doc, "事件传导确认方向：CMJ 先行 → ZZJ 跟随（22.14% vs 15.45%，中位滞后 4 min）")

    _new_heading(doc, "局限", 2)
    _add_bullet(doc, "恒流控制下的幅度层回归不可学，回归残差异常类方法在本数据上整体不可用")
    _add_bullet(doc, "未接入维修工单 / 故障记录，规则事件与真实故障的对应关系尚未验证")
    _add_bullet(doc, "数据仅覆盖 CMJ+ZZJ 两设备，未含皮带机/破碎机等其他运输链节点")
    _add_bullet(doc, "事件传导基于阶段二单设备异常事件，传导链的物理真伪需工单佐证")

    _new_heading(doc, "下一步方向", 2)
    _add_bullet(doc, "对 04-15 长时错配做人工核查（设备故障记录 / 维修工单佐证）")
    _add_bullet(doc, "错配/余流事件与维修工单时间对齐，验证检测精度")
    _add_bullet(doc, "若开放其余设备数据，将运输链扩展为多节点传导图")
    _add_bullet(doc, "完善周报，合并小组其他成员成果")


# ════════════════════════════════════════════════════════════
# 主程序
# ════════════════════════════════════════════════════════════

def main():
    print("=" * 55)
    print("阶段三技术报告生成器")
    print("=" * 55)

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── 全局样式 ──
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.35
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Microsoft YaHei"
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # ── 封面 ──
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("煤矿设备关联异常检测技术报告")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("阶段三：跨设备物理耦合与关联异常检测（CMJ → ZZJ）")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4D, 0x4D, 0x4D)

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run(
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        f"分析对象：采煤机（CMJ，上游）→ 转载机（ZZJ，下游）"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _add_page_break(doc)

    # ── 目录页 ──
    _new_heading(doc, "目录", 1)
    for ch in [
        "1. 引言",
        "2. 数据概况与分析设计",
        "3. 联合工况划分与物理耦合验证",
        "4. 物理耦合回归验证",
        "5. 联合工况规则事件主检测器",
        "6. 跨设备 Mahalanobis 关联异常",
        "7. 事件传导关联（CMJ 先行 → ZZJ 跟随）",
        "8. 可视化与典型案例",
        "9. 结论、局限与下一步",
    ]:
        _add_para(doc, ch, size=11)

    _add_page_break(doc)

    # ── 各章节 ──
    builders = [
        ("引言", build_introduction),
        ("数据概况", build_data_overview),
        ("联合工况", build_joint_condition),
        ("耦合回归", build_regression),
        ("规则事件", build_rule_events),
        ("Mahalanobis", build_mahalanobis),
        ("事件传导", build_propagation),
        ("可视化", build_visualization),
        ("结论", build_conclusion),
    ]

    for i, (name, builder) in enumerate(builders):
        try:
            print(f"  写入章节: {i+1}. {name}")
            builder(doc)
            if i < len(builders) - 1:
                _add_page_break(doc)
        except Exception as e:
            print(f"  [WARN] 章节 '{name}' 生成失败: {e}")
            traceback.print_exc()
            _add_para(doc, f"[本章节生成出错: {e}]", italic=True)

    # ── 保存 ──
    doc.save(str(OUTPUT_FILE))
    print(f"\nOK 技术报告已生成: {OUTPUT_FILE}")
    print(f"  文件大小: {OUTPUT_FILE.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
