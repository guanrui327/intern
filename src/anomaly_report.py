# -*- coding: utf-8 -*-
"""阶段二：报告生成模块。

生成 phase2_report.md 和 phase2_report.docx。

支持两种结果结构：
  - 嵌套结构（CMJ）：device_results["cmj"] = {"截割部_工况": {...}, ...}
  - 扁平结构（ZZJ）：device_results["zzj"] = {"merged_events": ..., ...}
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd


def _get_report_sections(
    device_results: dict[str, dict],
) -> list[tuple[str, dict, str]]:
    """将 device_results 统一为可迭代的 section 列表。

    每个元素: (section_title, result_dict, file_prefix)
    - section_title: 用于报告标题的显示文字（如"采煤机 - 截割部"）
    - result_dict:   该 section 的完整结果 dict
    - file_prefix:   输出文件前缀（如"cmj_截割部"、"zzj"）
    """
    sections: list[tuple[str, dict, str]] = []

    device_map = [("cmj", "采煤机"), ("zzj", "转载机")]

    for device_key, dev_name in device_map:
        result = device_results.get(device_key)
        if result is None:
            continue

        # ── 扁平结果：直接为一个 section ──
        if isinstance(result, dict) and any(
            k in result for k in ("merged_events", "mahal_events", "if_events")
        ):
            sections.append((dev_name, result, device_key))
            continue

        # ── 嵌套结果（分部位） ──
        if isinstance(result, dict):
            for part_cond_col, part_result in result.items():
                part_short = part_cond_col.replace("_工况", "")
                sections.append(
                    (f"{dev_name} - {part_short}", part_result, f"{device_key}_{part_short}")
                )

    return sections


def generate_markdown(
    device_results: dict[str, dict],
    output_path: str | Path,
) -> Path:
    """生成阶段二 Markdown 报告。"""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    sections = _get_report_sections(device_results)

    lines = [
        "# 阶段二：异常检测报告",
        "",
        f"> 生成时间：{pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}",
        "> 说明：基于阶段一工况划分 + 滑动窗口特征 + 多变量异常检测",
        "",
        "---",
        "## 总体概览",
        "",
    ]

    for title, res, _prefix in sections:
        events = res.get("merged_events")
        n_anom = res.get("n_anomaly_timepoints", 0)
        methods = res.get("methods", [])

        lines.append(f"### {title}")
        lines.append("")
        if events is not None and not events.empty:
            lines.append(f"- **总事件记录数**: {len(events)}")
            lines.append(f"- **异常时间点**: {n_anom}")
            lines.append(f"- **检测方法**: {', '.join(methods)}")
        else:
            lines.append("- 无异常检测结果")
        lines.append("")

    lines.extend([
        "---",
        "## 1. 分工况基线",
        "",
        "每种工况下各监测参数的统计基线（均值、中位数、IQR、p5/p95）。",
        "",
    ])

    for title, res, _prefix in sections:
        profiles = res.get("profiles", {})
        if not profiles:
            continue
        lines.append(f"### {title}")
        lines.append("")
        for cond_col, profile in profiles.items():
            if profile is None or profile.empty:
                continue
            n_params = profile["参数"].nunique()
            n_conds = profile["工况"].nunique()
            lines.append(f"- **{cond_col}**: {n_params} 参数 × {n_conds} 工况")
        lines.append("")

    lines.extend([
        "---",
        "## 2. 多变量异常检测",
        "",
    ])

    for title, res, _prefix in sections:
        mahal = res.get("mahal_events")
        iforest = res.get("if_events")
        residual = res.get("residual_events")

        lines.append(f"### {title}")
        lines.append("")

        if mahal is not None and not mahal.empty:
            n_anom = mahal["is_anomaly"].sum()
            lines.append(f"- **Mahalanobis 距离**: {n_anom}/{len(mahal)} 异常点")
            lines.append(f"  - 各工况 χ² 阈值自适应计算")
            lines.append(f"  - 特征贡献分解：对每个异常点输出 Top-3 贡献特征及百分比")
        else:
            lines.append("- **Mahalanobis 距离**: 未计算或无有效数据")

        if iforest is not None and not iforest.empty:
            n_anom = iforest["is_anomaly"].sum()
            lines.append(f"- **Isolation Forest**: {n_anom}/{len(iforest)} 异常点")
        else:
            lines.append("- **Isolation Forest**: 未计算或无有效数据")

        if residual is not None and not residual.empty:
            n_anom = residual["is_anomaly"].sum()
            lines.append(f"- **残差异常检测 (AR 前向预测)**: {n_anom}/{len(residual)} 异常点")
        else:
            lines.append("- **残差异常检测**: 未计算或无有效数据")

        lines.append("")

    lines.extend([
        "---",
        "## 3. 单变量异常检测（IQR + 3σ）",
        "",
    ])

    for title, res, _prefix in sections:
        value_anomalies = res.get("value_anomalies", {})
        if not value_anomalies:
            continue
        lines.append(f"### {title}")
        lines.append("")
        for cond_col, anomaly_df in value_anomalies.items():
            if anomaly_df is None or anomaly_df.empty:
                continue
            n_anom = anomaly_df["异常(短段过滤)"].sum()
            total = len(anomaly_df)
            pct = n_anom / total * 100 if total > 0 else 0
            severity_counts = anomaly_df[anomaly_df["异常(短段过滤)"]]["严重程度"].value_counts()
            lines.append(f"- **{cond_col}**: {n_anom}/{total} = {pct:.1f}% 异常点")
            if not severity_counts.empty:
                parts = [f"    - {k}: {v}" for k, v in severity_counts.items()]
                lines.extend(parts)
        lines.append("")

    lines.extend([
        "---",
        "## 4. 工况切换频率",
        "",
    ])

    for title, res, _prefix in sections:
        trans_rates = res.get("transition_rates", {})
        if not trans_rates:
            continue
        lines.append(f"### {title}")
        lines.append("")
        for cond_col, rate_df in trans_rates.items():
            if rate_df is None or rate_df.empty:
                continue
            lines.append(f"**{cond_col}**:")
            lines.append("")
            lines.append("| 工况 | 段数 | 总时长(min) | 平均段长(min) | 切换次数/小时 |")
            lines.append("|------|------|------------|-------------|-------------|")
            for _, r in rate_df.iterrows():
                lines.append(
                    f"| {r['工况']} | {r['段数']} | {r['总时长(min)']} | "
                    f"{r['平均段长(min)']} | {r['切换次数/小时']} |"
                )
            lines.append("")

    lines.extend([
        "---",
        "## 5. 图表输出",
        "",
    ])

    for title, _res, prefix in sections:
        lines.append(f"### {title}")
        lines.append("")
        lines.append(f"- [Mahalanobis 时间线](anomalies/{prefix}_mahalanobis_timeline.png)")
        lines.append(f"- [Mahalanobis 特征贡献分解](anomalies/{prefix}_feature_breakdown.png)")
        lines.append(f"- [Isolation Forest vs Mahalanobis 对比](anomalies/{prefix}_if_comparison.png)")
        lines.append(f"- [归因总结](anomalies/{prefix}_interpretation_summary.png)")
        lines.append(f"- [滑动窗口特征仪表板](anomalies/{prefix}_window_features.png)")
        lines.append("")

    lines.extend([
        "---",
        "## 6. 关键发现",
        "",
        "详细异常事件列表见 CSV 文件：",
        "",
    ])
    for _title, _res, prefix in sections:
        lines.append(f"- `anomalies/{prefix}_merged_events.csv`")

    lines.extend([
        "",
        "---",
        "## 7. 下一步建议",
        "",
        "1. 对持续异常时段进行根因追溯",
        "2. 结合设备维修记录验证异常检测准确性",
        "3. 建立在线 anomaly scoring 接口",
        "",
    ])

    content = "\n".join(lines)
    output_path.write_text(content, encoding="utf-8")
    print(f"  阶段二报告: {output_path}")
    return output_path


def build_docx_report(
    device_results: dict[str, dict],
    output_path: str | Path,
    image_dir: str | Path,
) -> Path:
    """生成阶段二 DOCX 报告。

    复用 generate_report_docx.py 的样式，插入图表 + 表格。
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    image_dir = Path(image_dir)

    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  [WARN] python-docx 未安装，跳过 DOCX 生成")
        # 保存为 .md 作为 fallback
        md_path = output_path.with_suffix(".md")
        generate_markdown(device_results, md_path)
        return md_path

    sections = _get_report_sections(device_results)

    doc = Document()

    # ── 封面标题 ──
    title = doc.add_heading("阶段二：异常检测报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"生成时间：{pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    # ── 总体概览 ──
    doc.add_heading("总体概览", level=1)
    for title_text, res, _prefix in sections:
        events = res.get("merged_events")
        n_anom = res.get("n_anomaly_timepoints", 0)
        methods = res.get("methods", [])

        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        if events is not None and not events.empty:
            p.add_run(f"总事件 {len(events)} 条, 异常时间点 {n_anom} 个, "
                      f"方法: {', '.join(methods)}")
        else:
            p.add_run("无异常检测结果")

    # ── 插入图片 ──
    doc.add_heading("异常检测图表", level=1)
    for title_text, _res, prefix in sections:
        doc.add_heading(title_text, level=2)
        image_patterns = [
            ("mahalanobis_timeline", "Mahalanobis 时间线"),
            ("feature_breakdown", "特征贡献分解"),
            ("if_comparison", "IF vs Mahalanobis 对比"),
            ("interpretation_summary", "归因总结"),
            ("window_features", "滑动窗口特征仪表板"),
        ]
        for fname_prefix, caption in image_patterns:
            img_path = image_dir / f"{prefix}_{fname_prefix}.png"
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(5.5))
                doc.add_paragraph(caption, style="Caption")

    # ── 异常事件 Top-10 表 ──
    doc.add_heading("Top-10 异常事件（按分数排序）", level=1)
    for title_text, res, _prefix in sections:
        events = res.get("merged_events")
        if events is None or events.empty:
            continue

        doc.add_heading(title_text, level=2)
        anom = events[events.get("any_anomaly", False)].copy()
        if anom.empty:
            doc.add_paragraph("无异常事件")
            continue

        # 按分数绝对值排序取 Top-10
        anom["abs_score"] = anom["分数"].abs()
        top = anom.sort_values("abs_score", ascending=False).head(10)

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "时间戳"
        hdr[1].text = "工况"
        hdr[2].text = "方法"
        hdr[3].text = "分数"
        hdr[4].text = "归因"

        for _, row in top.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row["时间戳"])[:16]
            cells[1].text = str(row.get("工况", ""))
            cells[2].text = str(row.get("方法", ""))
            cells[3].text = f"{row['分数']:.2f}"
            cells[4].text = str(row.get("interpretation", ""))[:60]

    doc.save(str(output_path))
    print(f"  阶段二 DOCX: {output_path}")
    return output_path
