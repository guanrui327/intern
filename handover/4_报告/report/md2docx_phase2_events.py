# -*- coding: utf-8 -*-
"""Markdown → DOCX：阶段二_异常事件逐事件分析.md。
基于 md2docx_workreview.py 的轻量转换器；图片路径优先相对 md 目录解析，
失败则回退到仓库根目录（报告内 ![](output/...) 相对仓库根）。"""
import os
import re
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = r"D:\cursorproject\intern\report\阶段二_异常事件逐事件分析.md"
DST = r"D:\cursorproject\intern\report\阶段二_异常事件逐事件分析.docx"
IMG_WIDTH_CM = 14.0

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
CODE_RE = re.compile(r"`([^`]+)`")
IMG_RE = re.compile(r"^\s*!\[([^\]]*)\]\((\S+)\)\s*$")
REPO_ROOT = r"D:\cursorproject\intern"


def add_runs(par, text):
    parts = BOLD_RE.split(text)
    for i, seg in enumerate(parts):
        if seg == "":
            continue
        if i % 2 == 1:
            r = par.add_run(seg)
            r.bold = True
        else:
            subs = CODE_RE.split(seg)
            for j, sub in enumerate(subs):
                if sub == "":
                    continue
                r = par.add_run(sub)
                if j % 2 == 1:
                    r.font.name = "Consolas"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")


def set_cn(run, size=10.5, bold=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if color:
        run.font.color.rgb = RGBColor(*color)


def resolve_image(path):
    md_dir = os.path.dirname(os.path.abspath(SRC))
    for base in (md_dir, REPO_ROOT):
        full = os.path.join(base, path)
        if os.path.exists(full):
            return full
    return None


def set_outline(p, level):
    """给段落设 Word 大纲级别（0=标题1），让 TOC 域能收集标题。"""
    pPr = p._p.get_or_add_pPr()
    el = pPr.find(qn("w:outlineLvl"))
    if el is None:
        el = OxmlElement("w:outlineLvl")
        pPr.append(el)
    el.set(qn("w:val"), str(level))


def add_toc_field(doc):
    """插入 Word 目录域：打开后右键→更新域，自动生成带页码目录。"""
    par = doc.add_paragraph()
    run = par.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    hint = OxmlElement("w:t")
    hint.text = "（在 Word 中选中此处，右键→更新域，生成目录）"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, hint, fld_end):
        run._r.append(el)


def enable_update_fields(doc):
    """打开文档时自动更新域（TOC 页码）。"""
    settings = doc.settings.element
    el = settings.find(qn("w:updateFields"))
    if el is None:
        el = OxmlElement("w:updateFields")
        settings.append(el)
    el.set(qn("w:val"), "true")


def add_picture_line(doc, path):
    full = resolve_image(path)
    if full is None:
        print(f"[warn] 图片不存在，跳过: {path}")
        return
    doc.add_picture(full, width=Cm(IMG_WIDTH_CM))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    lines = [l.rstrip("\n") for l in open(SRC, encoding="utf-8")]
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("---"):
            i += 1
            continue
        if not line.strip():
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:\-|]+\|?$", lines[i + 1]):
            header = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            ncols = len(header)
            tbl = doc.add_table(rows=1, cols=ncols)
            tbl.style = "Light Grid Accent 1"
            for j, h in enumerate(header):
                add_runs(tbl.rows[0].cells[j].paragraphs[0], h)
                for r in tbl.rows[0].cells[j].paragraphs[0].runs:
                    set_cn(r, size=10, bold=True)
            for row in rows:
                cells = tbl.add_row().cells
                for j in range(ncols):
                    if j < len(row):
                        add_runs(cells[j].paragraphs[0], row[j])
                    for r in cells[j].paragraphs[0].runs:
                        set_cn(r, size=10)
            doc.add_paragraph()
            continue

        m = IMG_RE.match(line)
        if m:
            add_picture_line(doc, m.group(2))
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            title_text = m.group(2)
            p = doc.add_paragraph()
            r = p.add_run(title_text)
            sizes = {1: 18, 2: 15, 3: 12.5, 4: 11.5, 5: 11, 6: 10.5}
            set_cn(r, size=sizes.get(level, 11), bold=True, color=(31, 78, 121) if level <= 3 else None)
            if title_text == "目录":
                # Word 原生目录域；跳过 md 里的静态目录列表避免重复
                add_toc_field(doc)
                i += 1
                while i < len(lines):
                    l = lines[i].strip()
                    if not l or re.match(r"^[-*]\s", l):
                        i += 1
                    else:
                        break
                continue
            # 标题设大纲级别，TOC 域才能收集（markdown #→大纲0）
            set_outline(p, level - 1)
            i += 1
            continue

        if line.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_runs(p, line.lstrip(">").strip())
            for r in p.runs:
                set_cn(r, size=10, color=(89, 89, 89))
            i += 1
            continue

        m = re.match(r"^\s*(?:[-*]|\d+\.)\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Bullet" if not re.match(r"^\s*\d", line) else "List Number")
            add_runs(p, m.group(1))
            for r in p.runs:
                set_cn(r, size=10.5)
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, line)
        for r in p.runs:
            set_cn(r, size=10.5)
        i += 1

    enable_update_fields(doc)
    doc.save(DST)
    print(f"OK -> {DST}")


if __name__ == "__main__":
    sys.exit(main())
