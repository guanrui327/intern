# -*- coding: utf-8 -*-
"""阶段二：异常检测技术报告生成器（按部位版）。

基于按部位工况独立运行的 pipeline 输出，为每个 CMJ 部位生成独立的
技术分析章节，ZZJ 保持单章节不变。

独立运行：python src/generate_tech_report.py
输出文件：output/phase2/phase2_tech_report.docx
"""

from __future__ import annotations

import os
import sys
import traceback
from datetime import datetime
from pathlib import Path

import numpy as np
import pandas as pd

# ── DOCX ──
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ════════════════════════════════════════════════════════════
# 路径
# ════════════════════════════════════════════════════════════
BASE = Path(__file__).resolve().parent.parent
PHASE2_DIR = BASE / "output" / "phase2"
ANOMALIES_DIR = PHASE2_DIR / "anomalies"
PROFILES_DIR = PHASE2_DIR / "profiles"
OUTPUT_FILE = PHASE2_DIR / "phase2_tech_report.docx"

# ── CMJ 部位定义 ──
CMJ_PARTS = [
    ("截割部", "截割部_工况"),
    ("牵引部", "牵引部_工况"),
    ("油泵",   "油泵_工况"),
    ("破碎机", "破碎机_工况"),
]

# 部位→监测参数关键词（来自 config.py CMJ_PART_MONITOR_MAP）
PART_MONITOR_KWS = {
    "截割部": ["滚筒", "电机_电流", "电机_温度", "角度"],
    "牵引部": ["牵引", "电机_电流", "电机_温度", "速度"],
    "油泵":   ["油泵", "电流", "温度", "油压"],
    "破碎机": ["破碎机", "电流", "温度"],
}

# 部位→工况数（来自实际 pipeline 输出）
PART_COND_COUNTS = {
    "截割部": 7,
    "牵引部": 4,
    "油泵":   3,
    "破碎机": 3,
}


# ════════════════════════════════════════════════════════════
# 工具函数
# ════════════════════════════════════════════════════════════

def _read_csv(path: Path) -> pd.DataFrame | None:
    """安全读取 CSV，utf-8-sig 编码。"""
    if not path.exists():
        return None
    try:
        return pd.read_csv(path, encoding="utf-8-sig")
    except Exception as e:
        print(f"  [WARN] 无法读取 {path.name}: {e}")
        return None


def _new_heading(doc: Document, text: str, level: int = 1):
    """添加标题。"""
    p = doc.add_heading(text, level=level)
    return p


def _add_para(doc: Document, text: str, bold: bool = False, italic: bool = False,
              size: int | None = None, spacing_after: int = 6):
    """添加段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = 1.35
    return p


def _add_bullet(doc: Document, text: str, level: int = 0):
    """添加要点列表。"""
    p = doc.add_paragraph(text, style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p


def _add_table(doc: Document, df: pd.DataFrame, caption: str = "", max_rows: int = 20):
    """将 DataFrame 渲染为表格。"""
    if df.empty:
        _add_para(doc, "（无数据）", italic=True)
        return None

    df = df.head(max_rows)
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    for i, (_, row) in enumerate(df.iterrows()):
        for j, col in enumerate(df.columns):
            cell = table.rows[i + 1].cells[j]
            val = row[col]
            cell.text = str(val) if not pd.isna(val) else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def _add_formula(doc: Document, latex: str):
    """添加公式文本（等宽字体居中）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(latex)
    run.font.name = "Consolas"
    run.italic = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def _add_image(doc: Document, path: Path, width_inches: float = 5.5, caption: str = ""):
    """插入图片+图题。"""
    if not path.exists():
        _add_para(doc, f"[图片不存在: {path.name}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width_inches))
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(f"图：{caption}")
        run.italic = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)


def _add_page_break(doc: Document):
    """分页。"""
    doc.add_page_break()


def _part_profile_path(part_name: str) -> Path:
    """返回部位基线 CSV 路径。"""
    return PROFILES_DIR / f"cmj_baseline_{part_name}_工况.csv"


def _part_anomaly_path(part_name: str, kind: str) -> Path:
    """返回部位异常检测 CSV 路径。"""
    return ANOMALIES_DIR / f"cmj_{part_name}_{kind}.csv"


def _part_img_path(part_name: str, img: str) -> Path:
    """返回部位图片路径。"""
    return ANOMALIES_DIR / f"cmj_{part_name}_{img}.png"


def _part_value_path(part_name: str) -> Path:
    """返回部位单变量异常 CSV 路径。"""
    return ANOMALIES_DIR / f"value_anomalies_{part_name}_工况.csv"


def _add_part_summary_table(doc: Document, caption: str, rows: list[dict]):
    """通用部位汇总表生成。"""
    if rows:
        _add_table(doc, pd.DataFrame(rows), caption)


# ════════════════════════════════════════════════════════════
# 章节生成器
# ════════════════════════════════════════════════════════════

def build_introduction(doc: Document):
    """第1章：引言"""
    _new_heading(doc, "引言", 1)

    _add_para(doc, (
        "本报告为「阶段二：异常检测」的技术总结文档。阶段二以阶段一输出的带工况宽表为输入，"
        "对采煤机（CMJ）和转载机（ZZJ）的运行数据进行多维度异常检测分析。"
    ))

    _add_para(doc, (
        "阶段一完成了设备分部位工况划分（CMJ 的截割部、牵引部、油泵、破碎机四个独立部位），"
        "参数相关性分析和数据质量核查，输出了带部位工况标签的宽表数据。"
        "阶段二在此基础上，按每个部位各自的工况列独立运行完整的异常检测 pipeline，"
        "使各部位的异常检测标准与其工况特点精准对齐。"
    ))

    _new_heading(doc, "分析目标", 2)
    _add_bullet(doc, "按设备部位建立分工况的正常运行基线（均值、中位数、IQR、p5/p95）")
    _add_bullet(doc, "通过滑动窗口提取时域特征（RMS、斜率、峰峰值），监测参数动态变化")
    _add_bullet(doc, "运用 Mahalanobis 距离检测跨参数联合异常")
    _add_bullet(doc, "运用 Isolation Forest 检测非线性多变量异常")
    _add_bullet(doc, "通过残差分析检测时间序列中的渐进式偏离")
    _add_bullet(doc, "通过单变量 IQR+3σ 检测参数级离群点")
    _add_bullet(doc, "多方法事件合并，生成统一的异常事件日志")

    _new_heading(doc, "分析范围", 2)
    _add_para(doc, (
        "采煤机（CMJ）：21 个监测参数，4 个独立部位工况列（截割部_工况、牵引部_工况、"
        "油泵_工况、破碎机_工况），每个部位按各自工况列独立建模。"
        "覆盖 2024 年 4–5 月共约 39,126 分钟运行数据。"
    ))
    _add_para(doc, (
        "转载机（ZZJ）：13 个监测参数，1 个工况列（工况），覆盖同期约 87,327 分钟运行数据。"
    ))

    _new_heading(doc, "检测方法矩阵", 2)
    method_df = pd.DataFrame({
        "方法": ["Mahalanobis 距离", "Isolation Forest", "残差分析（AR预测）", "单变量 IQR+3σ"],
        "类型": ["多变量", "多变量", "单变量*时序", "单变量"],
        "原理": [
            "稳健协方差估计 + χ² 阈值",
            "随机树异常隔离",
            "窗口均值预测 + Z-score",
            "四分位距 + 正态分布"],
        "输出": [
            "马氏距离、特征贡献占比",
            "异常分数、Top 偏离特征",
            "残差 Z-score、实际 vs 预测",
            "z-score、IQR/3σ 双阈值",
        ],
    })
    _add_table(doc, method_df, "表1：异常检测方法概览")
    _add_para(doc, "* 残差分析虽按参数逐列计算，但参数间是独立的，归为单变量类别。", italic=True, size=9)


def build_data_overview(doc: Document):
    """第2章：数据概况与分析设计"""
    _new_heading(doc, "数据概况与分析设计", 1)

    _new_heading(doc, "输入数据", 2)
    _add_para(doc, (
        "输入为阶段一输出的带工况宽表，1 分钟等间隔采样，DatetimeIndex 格式。"
    ))

    # Read per-part profiles for parameter counts
    part_params = {}
    for pname, pcond in CMJ_PARTS:
        dfp = _read_csv(_part_profile_path(pname))
        if dfp is not None:
            part_params[pname] = {
                "参数数": dfp["参数"].nunique(),
                "工况数": dfp["工况"].nunique(),
            }

    overview = pd.DataFrame({
        "设备/部位": ["CMJ-截割部", "CMJ-牵引部", "CMJ-油泵", "CMJ-破碎机", "ZZJ"],
        "工况列": ["截割部_工况", "牵引部_工况", "油泵_工况", "破碎机_工况", "工况"],
        "工况数": [
            part_params.get("截割部", {}).get("工况数", "—"),
            part_params.get("牵引部", {}).get("工况数", "—"),
            part_params.get("油泵", {}).get("工况数", "—"),
            part_params.get("破碎机", {}).get("工况数", "—"),
            "3",
        ],
        "监测参数": [
            part_params.get("截割部", {}).get("参数数", "—"),
            part_params.get("牵引部", {}).get("参数数", "—"),
            part_params.get("油泵", {}).get("参数数", "—"),
            part_params.get("破碎机", {}).get("参数数", "—"),
            "13",
        ],
        "数据点数": ["≈39,126"] * 4 + ["≈87,327"],
    })
    _add_table(doc, overview, "表2：输入数据概况（按部位）")
    _add_para(doc, (
        "注：各部位在同一时间片内有不同的工况标签。例如时间 t 时，"
        "截割部可能处于「割煤高位」、牵引部处于「重载牵引」、"
        "油泵处于「重载」、破碎机处于「空载运行」——"
        "这种「一设备多部位、各部位独立工况」的结构决定了异常检测必须按部位独立运行。"
    ), italic=True, size=9)

    _new_heading(doc, "核心设计原则", 2)

    _add_para(doc, "（一）分部位 + 分工况分析", bold=True)
    _add_para(doc, (
        "煤矿设备的运行状态随工况剧烈变化——割煤中电流是待机状态的数倍，"
        "而不同部位（截割部 vs 牵引部 vs 油泵）关注的参数和运行规律完全不同。"
        "若跨部位/跨工况统一分析，正常的高位割煤数据会被误判为异常，"
        "而待机时的微小波动又会淹没真正异常。"
        "因此，每个部位按自身工况列独立建模基线，是本分析的核心前提。"
    ))

    _add_para(doc, "（二）滑动窗口特征", bold=True)
    _add_para(doc, (
        "原始 1 分钟采样数据本身包含噪声和瞬时波动。"
        "滑动窗口（5 分钟窗口、1 分钟步长）提取 RMS（反映能量水平）、斜率（反映变化趋势）、"
        "峰峰值（反映波动幅度）等特征，在降噪的同时保留了信号的动态特征。"
    ))

    _add_para(doc, "（三）多方法协同", bold=True)
    _add_para(doc, (
        "没有一种异常检测方法能捕获所有类型的异常。Mahalanobis 擅长联合参数偏移，"
        "Isolation Forest 捕捉非线性异常，残差分析捕捉渐进式漂移，"
        "单变量方法定位具体异常参数。四种方法协同，事件合并得到综合判断。"
    ))

    _add_para(doc, "（四）可解释性优先", bold=True)
    _add_para(doc, (
        "不只输出二元标签，还提供每个异常事件的归因信息——特征贡献分解、"
        "偏离方向、实际 vs 预测值——使现场工程师能快速理解异常原因。"
    ))


def build_baseline(doc: Document):
    """第3章：分工况基线"""
    _new_heading(doc, "分工况基线建模", 1)

    _new_heading(doc, "原理", 2)
    _add_para(doc, (
        "基线（Baseline Profile）是对每种工况下各监测参数的统计分布建模，"
        "作为后续异常检测的参考标准。"
    ))
    _add_formula(doc, "对参数 x 在工况 c 下：Baseline_c = {mean, median, std, IQR, p5, p95}")
    _add_para(doc, (
        "选取中位数而非均值作为中心趋势指标——中位数对离群点更稳健，"
        "在工业数据（常有传感器尖峰和短暂中断）中更可靠。"
    ))

    _new_heading(doc, "为什么按部位+分工况建模", 2)
    _add_para(doc, (
        "不同部位关注的参数不同：截割部侧重滚筒电流和摇臂角度，"
        "牵引部侧重牵引电流和速度，油泵侧重油压和温度，破碎机侧重电流和温度。"
        "将部位参数与部位工况对齐，才能正确区分「正常的工况差异」和「真正的运行异常」。"
    ))

    _new_heading(doc, "最小样本要求", 2)
    _add_para(doc, (
        "每种工况至少需要 60 个样本（1 小时）方可构建有统计意义的基线。"
        "样本不足的工况予以跳过，避免基于少量数据的错误推断。"
    ))

    _new_heading(doc, "结果", 2)
    _add_para(doc, "按部位/工况列的基线概况：", bold=True)

    part_rows = []
    for pname, pcond in CMJ_PARTS:
        dfp = _read_csv(_part_profile_path(pname))
        if dfp is not None:
            part_rows.append({
                "部位": pname,
                "工况列": pcond,
                "工况数": dfp["工况"].nunique(),
                "参数数": dfp["参数"].nunique(),
            })
    zzj_baseline = _read_csv(PROFILES_DIR / "zzj_baseline_工况.csv")
    if zzj_baseline is not None:
        part_rows.append({
            "部位": "ZZJ",
            "工况列": "工况",
            "工况数": zzj_baseline["工况"].nunique(),
            "参数数": zzj_baseline["参数"].nunique(),
        })
    if part_rows:
        _add_table(doc, pd.DataFrame(part_rows), "表3：各部位/设备基线概况")

    # Profile sample
    first_profile = _part_profile_path(CMJ_PARTS[0][1].replace("_工况", ""))
    # Actually CMJ_PARTS[0] = ("截割部", "截割部_工况"), so name = "截割部"
    first_profile = _part_profile_path("截割部")
    df_prof = _read_csv(first_profile)
    if df_prof is not None:
        _add_para(doc, "以截割部为例，基线数据样例如下：")
        _add_table(doc, df_prof.head(8), "表4：截割部基线样例（前 8 行）")


def build_window_features(doc: Document):
    """第4章：滑动窗口特征提取"""
    _new_heading(doc, "滑动窗口特征提取", 1)

    _new_heading(doc, "原理", 2)
    _add_para(doc, (
        "滑动窗口特征提取将原始时间序列转换为具有统计意义的特征序列。"
        "对每个长度为 5 帧的窗口计算："
    ))

    _add_para(doc, "均方根（RMS）：", bold=True)
    _add_formula(doc, "RMS = sqrt((1/N) * Σ(x_i²)),   i ∈ [t-4, t]")
    _add_para(doc, "RMS 反映信号的能量水平。")

    _add_para(doc, "线性斜率（Slope）：", bold=True)
    _add_formula(doc, "Slope = np.polyfit(t, x, 1)[0]")
    _add_para(doc, "斜率捕捉参数的瞬时变化趋势，是早期异常的关键指标。")

    _add_para(doc, "峰峰值（Peak-to-Peak）：", bold=True)
    _add_formula(doc, "P2P = max(x_i) - min(x_i),   i ∈ [t-4, t]")
    _add_para(doc, "峰峰值反映窗口内的波动幅度。")

    _new_heading(doc, "结果", 2)

    _add_para(doc, "各部位窗口特征数据规模：", bold=True)
    win_rows = []
    for pname, _pcond in CMJ_PARTS:
        win_path = PHASE2_DIR / "windows" / f"cmj_{pname}_windows.csv"
        dfw = _read_csv(win_path)
        if dfw is not None:
            n_feat = len([c for c in dfw.columns if c not in ["时间戳", "工况"]])
            win_rows.append({
                "部位": f"CMJ-{pname}",
                "行数": len(dfw),
                "特征数": n_feat,
            })
    zzj_win = _read_csv(PHASE2_DIR / "windows" / "zzj_sliding_windows.csv")
    if zzj_win is not None:
        n_feat_zzj = len([c for c in zzj_win.columns if c not in ["时间戳", "工况"]])
        win_rows.append({"部位": "ZZJ", "行数": len(zzj_win), "特征数": n_feat_zzj})
    if win_rows:
        _add_table(doc, pd.DataFrame(win_rows), "表5：窗口特征数据规模")

    # Per-part window feature images
    _add_para(doc, "各部位窗口特征仪表板：")
    for pname, _ in CMJ_PARTS:
        _add_image(doc, _part_img_path(pname, "window_features"),
                   width_inches=5.2, caption=f"CMJ-{pname} 滑动窗口特征")
    _add_image(doc, ANOMALIES_DIR / "zzj_window_features.png",
               width_inches=5.2, caption="ZZJ 滑动窗口特征")


def _read_part_mahalanobis_summary(pname: str) -> dict | None:
    """读取部位 Mahalanobis 结果摘要。"""
    df = _read_csv(_part_anomaly_path(pname, "mahalanobis"))
    if df is None or df.empty:
        return None
    return {
        "总样本": len(df),
        "异常数": int(df["is_anomaly"].sum()),
        "异常率": f"{df['is_anomaly'].sum() / len(df) * 100:.1f}%",
    }


def build_mahalanobis(doc: Document):
    """第5章：Mahalanobis 距离异常检测（按部位）"""
    _new_heading(doc, "Mahalanobis 距离异常检测", 1)

    _new_heading(doc, "原理", 2)
    _add_para(doc, (
        "马氏距离（Mahalanobis Distance）是一种多变量距离度量，考虑了变量之间的相关性。"
    ))
    _add_formula(doc, "D_M(x) = sqrt((x - μ)^T · Σ^(-1) · (x - μ))")
    _add_para(doc, (
        "当数据服从多元正态分布时，D_M² 服从 χ² 分布，自由度为特征维度，"
        "因此用 χ² 分位数作为判别阈值。使用 MinCovDet（MCD）稳健估计协方差矩阵，"
        "抵抗离群点干扰。"
    ))

    _new_heading(doc, "特征贡献分解", 2)
    _add_formula(doc, "贡献_j = (x_j - μ_j)² · Σ^(-1)_jj")
    _add_para(doc, (
        "选取贡献最大的 Top-3 特征作为异常归因。"
    ))

    _new_heading(doc, "结果（按部位）", 2)

    # Per-part summary table
    maha_rows = []
    for pname, _ in CMJ_PARTS:
        s = _read_part_mahalanobis_summary(pname)
        if s:
            maha_rows.append({"部位": pname, **s})
    zzj_mahal = _read_csv(ANOMALIES_DIR / "zzj_mahalanobis.csv")
    if zzj_mahal is not None:
        maha_rows.append({
            "部位": "ZZJ",
            "总样本": len(zzj_mahal),
            "异常数": int(zzj_mahal["is_anomaly"].sum()),
            "异常率": f"{zzj_mahal['is_anomaly'].sum() / len(zzj_mahal) * 100:.1f}%",
        })
    if maha_rows:
        _add_table(doc, pd.DataFrame(maha_rows), "表6：Mahalanobis 异常检测结果汇总")

    # Per-part detail: breakdown by condition
    for pname, pcond in CMJ_PARTS:
        dfm = _read_csv(_part_anomaly_path(pname, "mahalanobis"))
        if dfm is None:
            continue
        _add_para(doc, f"CMJ-{pname} 按工况明细：", bold=True, size=10)
        summary = dfm.groupby("工况").agg(
            总样本=("is_anomaly", "count"),
            异常数=("is_anomaly", "sum"),
            异常率=("is_anomaly", lambda x: f"{x.sum() / len(x) * 100:.1f}%"),
        ).reset_index()
        summary.columns = ["工况", "总样本", "异常数", "异常率"]
        _add_table(doc, summary, f"CMJ-{pname} Mahalanobis 工况级结果", max_rows=10)

        _add_image(doc, _part_img_path(pname, "mahalanobis_timeline"),
                   width_inches=5.5, caption=f"CMJ-{pname} Mahalanobis 时间线")
        _add_image(doc, _part_img_path(pname, "feature_breakdown"),
                   width_inches=5.2, caption=f"CMJ-{pname} 特征贡献分解")

    # ZZJ
    if zzj_mahal is not None:
        _add_para(doc, "ZZJ Mahalanobis 按工况明细：", bold=True, size=10)
        summary = zzj_mahal.groupby("工况").agg(
            总样本=("is_anomaly", "count"),
            异常数=("is_anomaly", "sum"),
            异常率=("is_anomaly", lambda x: f"{x.sum() / len(x) * 100:.1f}%"),
        ).reset_index()
        summary.columns = ["工况", "总样本", "异常数", "异常率"]
        _add_table(doc, summary, "ZZJ Mahalanobis 工况级结果", max_rows=10)
        _add_image(doc, ANOMALIES_DIR / "zzj_mahalanobis_timeline.png",
                   width_inches=5.5, caption="ZZJ Mahalanobis 时间线")
        _add_image(doc, ANOMALIES_DIR / "zzj_feature_breakdown.png",
                   width_inches=5.2, caption="ZZJ 特征贡献分解")


def build_isolation_forest(doc: Document):
    """第6章：Isolation Forest 异常检测（按部位）"""
    _new_heading(doc, "Isolation Forest 异常检测", 1)

    _new_heading(doc, "原理", 2)
    _add_para(doc, (
        "Isolation Forest 是一种基于树集成的无监督异常检测算法。"
        "异常点容易被孤立（位于稀疏区域），随机切分时路径长度更短。"
    ))
    _add_formula(doc, "s(x, N) = 2^(-E[h(x)] / c(N))")
    _add_para(doc, "s(x) 接近 1 表示高度可能是异常，接近 0 表示正常。")

    _new_heading(doc, "结果（按部位）", 2)

    if_rows = []
    for pname, _ in CMJ_PARTS:
        dfi = _read_csv(_part_anomaly_path(pname, "iforest"))
        if dfi is not None:
            n_anom = dfi["is_anomaly"].sum()
            if_rows.append({
                "部位": pname,
                "总样本": len(dfi),
                "异常数": int(n_anom),
                "异常率": f"{n_anom / len(dfi) * 100:.2f}%",
            })
    zzj_if = _read_csv(ANOMALIES_DIR / "zzj_iforest.csv")
    if zzj_if is not None:
        n_anom = zzj_if["is_anomaly"].sum()
        if_rows.append({
            "部位": "ZZJ",
            "总样本": len(zzj_if),
            "异常数": int(n_anom),
            "异常率": f"{n_anom / len(zzj_if) * 100:.2f}%",
        })
    if if_rows:
        _add_table(doc, pd.DataFrame(if_rows), "表7：Isolation Forest 结果汇总")

    # IF vs Mahalanobis comparison images per part
    for pname, _ in CMJ_PARTS:
        _add_image(doc, _part_img_path(pname, "if_comparison"),
                   width_inches=5.2, caption=f"CMJ-{pname} IF vs Mahalanobis 对比")
    _add_image(doc, ANOMALIES_DIR / "zzj_if_comparison.png",
               width_inches=5.2, caption="ZZJ IF vs Mahalanobis 对比")


def build_residual(doc: Document):
    """第7章：残差异常检测（按部位）"""
    _new_heading(doc, "残差异常检测（AR 前向预测）", 1)

    _new_heading(doc, "原理", 2)
    _add_formula(doc, "pred(t) = mean(x[t-5], x[t-4], ..., x[t-1])   |   同工况下")
    _add_formula(doc, "residual(t) = x(t) - pred(t)")
    _add_formula(doc, "Z(t) = residual(t) / σ_residual")
    _add_para(doc, "残差 Z-score 超过 3.0 标记为异常。")

    _new_heading(doc, "结果（按部位）", 2)

    res_rows = []
    for pname, _ in CMJ_PARTS:
        dfr = _read_csv(_part_anomaly_path(pname, "residual"))
        if dfr is not None:
            n_anom = dfr["is_anomaly"].sum()
            zs = dfr["z_residual"]
            res_rows.append({
                "部位": pname,
                "总样本": len(dfr),
                "异常数": int(n_anom),
                "异常率": f"{n_anom / len(dfr) * 100:.1f}%",
                "Z中位数": f"{zs.median():.2f}",
            })
    zzj_res = _read_csv(ANOMALIES_DIR / "zzj_residual.csv")
    if zzj_res is not None:
        n_anom = zzj_res["is_anomaly"].sum()
        zs = zzj_res["z_residual"]
        res_rows.append({
            "部位": "ZZJ",
            "总样本": len(zzj_res),
            "异常数": int(n_anom),
            "异常率": f"{n_anom / len(zzj_res) * 100:.1f}%",
            "Z中位数": f"{zs.median():.2f}",
        })
    if res_rows:
        _add_table(doc, pd.DataFrame(res_rows), "表8：残差异常检测结果汇总")


def build_univariate(doc: Document):
    """第8章：单变量异常检测（按部位）"""
    _new_heading(doc, "单变量异常检测（IQR + 3σ）", 1)

    _new_heading(doc, "原理", 2)
    _add_para(doc, "使用两种经典统计方法，任一标记为异常即视为异常：")
    _add_formula(doc, "IQR 法: x < Q1 - 1.5*IQR  或   x > Q3 + 1.5*IQR")
    _add_formula(doc, "3σ 法: x < mean - 3*std  或   x > mean + 3*std")

    _new_heading(doc, "短段过滤", 2)
    _add_para(doc, (
        "连续异常至少需要 3 个样本（3 分钟）才判定为真实异常，"
        "短于 3 点的孤立异常视为噪声过滤掉。"
    ))

    _new_heading(doc, "结果（按部位）", 2)

    uv_rows = []
    for pname, _ in CMJ_PARTS:
        dfv = _read_csv(_part_value_path(pname))
        if dfv is not None and "异常(短段过滤)" in dfv.columns:
            n_anom = dfv["异常(短段过滤)"].sum()
            uv_rows.append({
                "部位": pname,
                "总样本": len(dfv),
                "异常数（过滤后）": int(n_anom),
                "异常率": f"{n_anom / len(dfv) * 100:.1f}%",
            })
    zzj_val = _read_csv(ANOMALIES_DIR / "value_anomalies_工况.csv")
    if zzj_val is not None and "异常(短段过滤)" in zzj_val.columns:
        n_anom = zzj_val["异常(短段过滤)"].sum()
        uv_rows.append({
            "部位": "ZZJ",
            "总样本": len(zzj_val),
            "异常数（过滤后）": int(n_anom),
            "异常率": f"{n_anom / len(zzj_val) * 100:.1f}%",
        })
    if uv_rows:
        _add_table(doc, pd.DataFrame(uv_rows), "表9：单变量异常检测结果汇总")

    # Per-part value anomaly images
    for pname, _ in CMJ_PARTS:
        img = ANOMALIES_DIR / f"value_anomalies_{pname}_工况.png"
        _add_image(doc, img, width_inches=5.2, caption=f"CMJ-{pname} 单变量异常时序图")
    _add_image(doc, ANOMALIES_DIR / "value_anomalies_工况.png",
               width_inches=5.2, caption="ZZJ 单变量异常时序图")


def build_merge(doc: Document):
    """第9章：多方法事件合并与归因分析（按部位）"""
    _new_heading(doc, "多方法事件合并与归因分析", 1)

    _new_heading(doc, "原理", 2)
    _add_para(doc, (
        "四种异常检测方法各有侧重，事件合并将各方法的检测结果汇总为统一的异常事件表。"
        "各方法事件通过时间戳对齐，any_anomaly=True 表示该时间片至少被一种方法标记为异常。"
        "interpretation 字段提供可读的归因文本。"
    ))

    _new_heading(doc, "结果（按部位）", 2)

    merge_rows = []
    for pname, _ in CMJ_PARTS:
        dfm = _read_csv(_part_anomaly_path(pname, "merged_events"))
        if dfm is not None and "any_anomaly" in dfm.columns:
            n_anom = dfm["any_anomaly"].sum()
            merge_rows.append({
                "部位": pname,
                "总记录": len(dfm),
                "异常时间点": int(n_anom),
            })
    zzj_merge = _read_csv(ANOMALIES_DIR / "zzj_merged_events.csv")
    if zzj_merge is not None and "any_anomaly" in zzj_merge.columns:
        n_anom = zzj_merge["any_anomaly"].sum()
        merge_rows.append({
            "部位": "ZZJ",
            "总记录": len(zzj_merge),
            "异常时间点": int(n_anom),
        })
    if merge_rows:
        _add_table(doc, pd.DataFrame(merge_rows), "表10：合并事件汇总")

    # Per-part interpretation images
    for pname, _ in CMJ_PARTS:
        _add_image(doc, _part_img_path(pname, "interpretation_summary"),
                   width_inches=5.2, caption=f"CMJ-{pname} 异常归因总结")
        # Sample events
        dfm = _read_csv(_part_anomaly_path(pname, "merged_events"))
        if dfm is not None and "interpretation" in dfm.columns and "any_anomaly" in dfm.columns:
            top_m = dfm[dfm["any_anomaly"]].head(5)
            display_cols = [c for c in ["时间戳", "工况", "方法", "interpretation"] if c in top_m.columns]
            if display_cols:
                _add_table(doc, top_m[display_cols], f"CMJ-{pname} 异常事件样例")

    _add_image(doc, ANOMALIES_DIR / "zzj_interpretation_summary.png",
               width_inches=5.2, caption="ZZJ 异常归因总结")
    if zzj_merge is not None and "interpretation" in zzj_merge.columns and "any_anomaly" in zzj_merge.columns:
        top_z = zzj_merge[zzj_merge["any_anomaly"]].head(5)
        display_cols = [c for c in ["时间戳", "工况", "方法", "interpretation"] if c in top_z.columns]
        if display_cols:
            _add_table(doc, top_z[display_cols], "ZZJ 异常事件样例")


def build_transition(doc: Document):
    """第10章：工况切换频率分析（按部位）"""
    _new_heading(doc, "工况切换频率分析", 1)

    _new_heading(doc, "重要性", 2)
    _add_para(doc, (
        "频繁的工况切换本身是一种运行特征。异常高的切换频率可能指示控制系统异常，"
        "过长的单工况持续时间可能表明设备卡滞在某一状态。"
    ))

    _new_heading(doc, "结果（按部位）", 2)

    for pname, pcond in CMJ_PARTS:
        tr_path = ANOMALIES_DIR / f"cmj_transition_rate_{pcond}.csv"
        df_tr = _read_csv(tr_path)
        if df_tr is None or df_tr.empty:
            continue
        _add_para(doc, f"CMJ-{pname} 工况切换频率：", bold=True)
        _add_table(doc, df_tr, f"CMJ-{pname} 工况切换")

    zzj_tr = _read_csv(ANOMALIES_DIR / "zzj_transition_rate_工况.csv")
    if zzj_tr is not None:
        _add_para(doc, "ZZJ 工况切换频率：", bold=True)
        _add_table(doc, zzj_tr, "ZZJ 工况切换")


def build_conclusion(doc: Document):
    """第11章：结论与建议"""
    _new_heading(doc, "结论与后续建议", 1)

    _new_heading(doc, "主要结论", 2)
    _add_para(doc, (
        "1. 分部位+分工况基线建模是本分析的核心前提——不同部位关注的参数不同，"
        "各部位在不同工况下的参数分布差异显著，跨部位/跨工况统一分析会产生严重的误分类。"
    ))
    _add_para(doc, (
        "2. 四种异常检测方法互为补充。各部位 Mahalanobis 异常率约 15-23%，"
        "反映了多参数联合偏移的普遍性；Isolation Forest 异常率约 0.1%（最极端的离群点）；"
        "残差分析约 1.5-1.8%（渐进式漂移）；单变量 IQR+3σ 约 1.1-2.5%（参数级离群）。"
    ))
    _add_para(doc, (
        "3. 特征贡献分解提供了可解释的异常归因。"
    ))

    _new_heading(doc, "注意事项与局限", 2)
    _add_para(doc, (
        "1. Mahalanobis 异常率偏高（15-23%），可能因数据分布偏离多元正态假设，"
        "或 χ² 阈值（α=0.001）过于宽松。建议后续调优至 α=0.0001。"
    ))
    _add_para(doc, (
        "2. 当前分析基于 1 分钟采样数据，秒级瞬变事件无法捕捉。"
    ))
    _add_para(doc, (
        "3. 未接入维修记录工单，无法验证异常与实际故障的对应关系。"
    ))

    _new_heading(doc, "下一步方向", 2)
    _add_bullet(doc, "调优参数：收紧 Mahalanobis α 阈值，降低异常率至 5% 以下")
    _add_bullet(doc, "在线接口：将基线模型部署为实时 anomaly scoring 服务")
    _add_bullet(doc, "根因追溯：对持续异常时段，结合维修记录做根因分析")
    _add_bullet(doc, "跨部位联动：分析 CMJ 各部位异常事件的时序相关性")
    _add_bullet(doc, "跨设备联动：分析采煤机与转载机异常事件的时序关联")


# ════════════════════════════════════════════════════════════
# 主程序
# ════════════════════════════════════════════════════════════

def main():
    print("=" * 55)
    print("阶段二技术报告生成器（按部位版）")
    print("=" * 55)

    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── 全局样式 ──
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.35
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Microsoft YaHei"
        hs.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # ── 封面 ──
    for _ in range(6):
        doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("煤矿设备异常检测技术报告")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("阶段二：基于按部位分工况的多维度异常检测")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x4D, 0x4D, 0x4D)

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run(
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        f"分析对象：采煤机（CMJ 4 部位）& 转载机（ZZJ）"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _add_page_break(doc)

    # ── 目录页 ──
    _new_heading(doc, "目录", 1)
    chapters = [
        "1. 引言",
        "2. 数据概况与分析设计",
        "3. 分工况基线建模",
        "4. 滑动窗口特征提取",
        "5. Mahalanobis 距离异常检测",
        "6. Isolation Forest 异常检测",
        "7. 残差异常检测（AR 前向预测）",
        "8. 单变量异常检测（IQR + 3σ）",
        "9. 多方法事件合并与归因分析",
        "10. 工况切换频率分析",
        "11. 结论与后续建议",
    ]
    for ch in chapters:
        _add_para(doc, ch, size=11)

    _add_page_break(doc)

    # ── 各章节 ──
    builders = [
        ("引言", build_introduction),
        ("数据概况", build_data_overview),
        ("基线", build_baseline),
        ("窗口特征", build_window_features),
        ("Mahalanobis", build_mahalanobis),
        ("Isolation Forest", build_isolation_forest),
        ("残差", build_residual),
        ("单变量", build_univariate),
        ("事件合并", build_merge),
        ("工况切换", build_transition),
        ("结论", build_conclusion),
    ]

    for i, (name, builder) in enumerate(builders):
        try:
            print(f"  写入章节: {i+1}. {name}")
            builder(doc)
            if i < len(builders) - 1:
                _add_page_break(doc)
        except Exception as e:
            print(f"  [WARN] 章节 '{name}' 生成失败: {e}")
            traceback.print_exc()
            _add_para(doc, f"[本章节生成出错: {e}]", italic=True)

    # ── 保存 ──
    doc.save(str(OUTPUT_FILE))
    print(f"\nOK 技术报告已生成: {OUTPUT_FILE}")
    print(f"  文件大小: {OUTPUT_FILE.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
