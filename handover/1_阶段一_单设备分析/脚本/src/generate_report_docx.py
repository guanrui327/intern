# -*- coding: utf-8 -*-
"""
阶段一：单设备分析报告 — DOCX 生成（排版优化版）
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ── 路径 ──
BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PHASE1 = os.path.join(BASE, 'output', 'phase1')
OUTPUT_FILE = os.path.join(PHASE1, '阶段一_单设备分析报告_v2.docx')

CHART_MAP = {
    'cmj_device_timeline':  ('采煤机设备级工况时间线',     'cmj_device_condition_timeline.png'),
    'cmj_device_pie':       ('采煤机设备级工况占比饼图',   'cmj_device_condition_pie.png'),
    'cmj_cutting_timeline': ('采煤机截割部工况时间线',     'cmj_截割部_timeline.png'),
    'cmj_traction_timeline':('采煤机牵引部工况时间线',     'cmj_牵引部_timeline.png'),
    'cmj_pump_timeline':    ('采煤机油泵工况时间线',        'cmj_油泵_timeline.png'),
    'cmj_crusher_timeline': ('采煤机破碎机工况时间线',     'cmj_破碎机_timeline.png'),
    'cmj_cut_current':      ('截割部电流箱线图',           'cmj_截割部_cut_current_boxplot.png'),
    'cmj_trac_current':     ('牵引部电流箱线图',           'cmj_牵引部_traction_current_boxplot.png'),
    'cmj_pump_current':     ('油泵电流箱线图',             'cmj_油泵_pump_current_boxplot.png'),
    'cmj_crusher_current':  ('破碎机电流箱线图',           'cmj_破碎机_crusher_current_boxplot.png'),
    'cmj_corr':             ('采煤机参数 Spearman 热力图', 'cmj_corr_heatmap.png'),
    'zzj_timeline':         ('转载机工况时间线',            'zzj_condition_timeline.png'),
    'zzj_pie':              ('转载机工况占比饼图',          'zzj_condition_pie.png'),
    'zzj_current':          ('转载机分工况电流箱线图',      'zzj_current_by_cond.png'),
    'zzj_corr':             ('转载机参数 Spearman 热力图',  'zzj_corr_heatmap.png'),
    'gap_detection':        ('数据空洞热力图',              'gap_detection.png'),
    'param_hierarchy':      ('参数层级图谱',                'param_hierarchy.png'),
    'cmj_profile':          ('跨工况参数 Profile 对比',    'cmj_all_params_by_device_profile.png'),
    'cluster_vs_截割部':    ('截割部聚类 vs 规则工况对比', 'cluster_vs_截割部.png'),
    'cluster_vs_牵引部':    ('牵引部聚类 vs 规则工况对比', 'cluster_vs_牵引部.png'),
    'cluster_vs_油泵':      ('油泵聚类 vs 规则工况对比',   'cluster_vs_油泵.png'),
    'cluster_vs_破碎机':    ('破碎机聚类 vs 规则工况对比', 'cluster_vs_破碎机.png'),
}

# ── 第二周报告图表映射 ──
W2_CHART_MAP = {
    'kw_截割部':       ('截割部 Kruskal-Wallis 热力图',   'kruskal_截割部_heatmap.png'),
    'kw_牵引部':       ('牵引部 Kruskal-Wallis 热力图',   'kruskal_牵引部_heatmap.png'),
    'kw_油泵':         ('油泵 Kruskal-Wallis 热力图',     'kruskal_油泵_heatmap.png'),
    'kw_破碎机':       ('破碎机 Kruskal-Wallis 热力图',   'kruskal_破碎机_heatmap.png'),
    'seg_截割部':      ('截割部段持续时间箱线图',         'segment_duration_截割部_工况.png'),
    'seg_牵引部':      ('牵引部段持续时间箱线图',         'segment_duration_牵引部_工况.png'),
    'seg_油泵':        ('油泵段持续时间箱线图',           'segment_duration_油泵_工况.png'),
    'seg_破碎机':      ('破碎机段持续时间箱线图',         'segment_duration_破碎机_工况.png'),
    'anomaly_截割部':  ('截割部异常段检测',               'anomalous_segments_截割部_工况.png'),
    'anomaly_牵引部':  ('牵引部异常段检测',               'anomalous_segments_牵引部_工况.png'),
    'anomaly_油泵':    ('油泵异常段检测',                 'anomalous_segments_油泵_工况.png'),
    'anomaly_破碎机':  ('破碎机异常段检测',               'anomalous_segments_破碎机_工况.png'),
    'value_anomaly_截割部': ('截割部值异常检测（时序）', 'value_anomalies_截割部.png'),
    'value_anomaly_牵引部': ('牵引部值异常检测（时序）', 'value_anomalies_牵引部.png'),
    'value_anomaly_油泵':   ('油泵值异常检测（时序）',   'value_anomalies_油泵.png'),
    'value_anomaly_破碎机': ('破碎机值异常检测（时序）', 'value_anomalies_破碎机.png'),
    'cluster_截割部':  ('截割部聚类 vs 规则工况',         'cluster_vs_截割部.png'),
    'cluster_牵引部':  ('牵引部聚类 vs 规则工况',         'cluster_vs_牵引部.png'),
    'cluster_油泵':    ('油泵聚类 vs 规则工况',           'cluster_vs_油泵.png'),
    'cluster_破碎机':  ('破碎机聚类 vs 规则工况',         'cluster_vs_破碎机.png'),
    'fi_截割部':       ('截割部特征重要性（Random Forest）','feature_importance_截割部.png'),
    'fi_牵引部':       ('牵引部特征重要性（Random Forest）','feature_importance_牵引部.png'),
    'fi_油泵':         ('油泵特征重要性（Random Forest）',  'feature_importance_油泵.png'),
    'fi_破碎机':       ('破碎机特征重要性（Random Forest）','feature_importance_破碎机.png'),
    'profile_电流':    ('切换聚合剖面 — 电流',             'transition_profile_device_电流.png'),
    'multi_电流速度_0':('切换指纹 — 电流+速度 #0',         'transition_multi_param_电流速度_0.png'),
    'multi_电流速度_1':('切换指纹 — 电流+速度 #1',         'transition_multi_param_电流速度_1.png'),
    'multi_电流速度_2':('切换指纹 — 电流+速度 #2',         'transition_multi_param_电流速度_2.png'),
    'multi_电流温度_0':('切换指纹 — 电流+温度 #0',         'transition_multi_param_电流温度_0.png'),
    'multi_电流温度_1':('切换指纹 — 电流+温度 #1',         'transition_multi_param_电流温度_1.png'),
    'multi_电流温度_2':('切换指纹 — 电流+温度 #2',         'transition_multi_param_电流温度_2.png'),
    'multi_全关键参数_0':('切换指纹 — 全关键参数 #0',      'transition_multi_param_全关键参数_0.png'),
    'multi_全关键参数_1':('切换指纹 — 全关键参数 #1',      'transition_multi_param_全关键参数_1.png'),
    'multi_全关键参数_2':('切换指纹 — 全关键参数 #2',      'transition_multi_param_全关键参数_2.png'),
    'profile_温度':    ('切换聚合剖面 — 温度',             'transition_profile_device_温度.png'),
    'profile_速度':    ('切换聚合剖面 — 速度',             'transition_profile_device_速度.png'),
    # ── 转载机 ──
    'kw_转载机':          ('转载机 Kruskal-Wallis 热力图',         'kruskal_转载机_heatmap.png'),
    'seg_转载机':         ('转载机段持续时间箱线图',               'segment_duration_工况.png'),
    'anomaly_转载机':     ('转载机异常段检测',                    'anomalous_segments_工况.png'),
    'cluster_转载机':     ('转载机聚类 vs 规则工况',              'cluster_vs_转载机.png'),
    'fi_转载机':          ('转载机特征重要性（Random Forest）',    'feature_importance_转载机.png'),
    'transition_zzj':     ('转载机工况转换分析（电流）',           'transition_工况_current.png'),
    # ── 分部位滞后互相关 ──
    'lag_cmj_截割部':     ('截割部电流 vs 温度滞后互相关',        'cmj_lagged_corr_截割部.png'),
    'lag_cmj_牵引部':     ('牵引部电流 vs 速度滞后互相关',        'cmj_lagged_corr_牵引部.png'),
    'lag_cmj_油泵':       ('油泵电流 vs 油压滞后互相关',          'cmj_lagged_corr_油泵.png'),
    'lag_cmj_破碎机':     ('破碎机电流 vs 温度滞后互相关',        'cmj_lagged_corr_破碎机.png'),
    'lag_zzj':            ('转载机电流 vs 转速滞后互相关',        'zzj_lagged_corr_电流_vs_转速.png'),
}

LQ, RQ = '“', '”'


# ── XML helpers ──
def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', 'D9D9D9'))
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_page_number(doc):
    """Add page number in footer: 第 X 页 / 共 Y 页"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.size = Pt(9)
    p.style.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    run = p.add_run('第 ')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # PAGE field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run1 = p.add_run()
    run1._r.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2 = p.add_run()
    run2._r.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3 = p.add_run()
    run3._r.append(fldChar2)

    run = p.add_run(' 页 / 共 ')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # NUMPAGES field
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run4 = p.add_run()
    run4._r.append(fldChar3)
    instrText2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    run5 = p.add_run()
    run5._r.append(instrText2)
    fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run6 = p.add_run()
    run6._r.append(fldChar4)

    run = p.add_run(' 页')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ── Table ──
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # remove default paragraph spacing inside cells
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0

    # header
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2F5496')
        set_cell_border(cell,
            top={'val': 'single', 'sz': '6', 'color': '2F5496'},
            bottom={'val': 'single', 'sz': '8', 'color': '1F3864'},
        )

    # rows, alternating
    for ri, rdata in enumerate(rows):
        row = table.add_row()
        for ci, text in enumerate(rdata):
            cell = row.cells[ci]
            cell.text = str(text)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
            # alternating row color
            if ri % 2 == 0:
                set_cell_shading(cell, 'F2F2F2')

    # column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    return table


# ── Spacer ──
def spacer(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run('')
    run.font.size = Pt(pt)


def add_caption(doc, caption_text):
    """添加居中的图表/表格标题（粗体、10pt）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(caption_text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


# ═══════════════════════════════════════════
def build_report():
    doc = Document()

    # ── 页面设置（A4） ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ── 全局样式 ──
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Microsoft YaHei'
        hs.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        if level == 1:
            hs.font.size = Pt(18)
            hs.paragraph_format.space_before = Pt(18)
            hs.paragraph_format.space_after = Pt(10)
        elif level == 2:
            hs.font.size = Pt(14)
            hs.paragraph_format.space_before = Pt(14)
            hs.paragraph_format.space_after = Pt(6)
        else:
            hs.font.size = Pt(12)
            hs.paragraph_format.space_before = Pt(10)
            hs.paragraph_format.space_after = Pt(4)

    # ── 页脚页码 ──
    add_page_number(doc)

    # ═══ 封面 ═══
    for _ in range(5):
        spacer(doc, 14)

    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run('阶段一：单设备分析报告')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    # 分割线
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run('━' * 32)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    spacer(doc, 4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('采煤机（CMJ）+ 转载机（ZZJ）')
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run('工况划分与统计分析')
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    spacer(doc, 16)

    info_lines = [
        ('数据来源', '大海则煤矿  2024-04-01 ~ 2024-06-01'),
        ('数据格式', 'On-change 存储传感器数据'),
        ('分析设备', '采煤机（CMJ）、转载机（ZZJ）'),
        ('阶段工期', '2026-07-01 ~ 2026-07-15'),
        ('报告日期', '2026-07-07'),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'{label}：')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run = p.add_run(value)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_page_break()

    # ═══ 目录 ═══
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 数据概览',
        '2. 数据预处理',
        '3. 数据空洞检测',
        '4. 参数层级图谱与分部位工况',
        '5. 工况划分',
        '6. 分工况统计分析',
        '7. 可视化结果',
        '8. 工况转换分析',
        '9. 聚类验证规则工况',
        '10. 参数关联深入与跨工况对比',
        '11. 关键发现',
        '12. 阶段二预备',
        '附录：图表清单',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(14), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.bold = True

    doc.add_page_break()

    # ═══ 1. 数据概览 ═══
    doc.add_heading('1. 数据概览', level=1)

    doc.add_heading('1.1 数据来源', level=2)
    doc.add_paragraph(
        '本项目使用大海则煤矿 2024-04-01 至 2024-06-01 的传感器实时数据。'
        '数据以 on-change 格式存储，即仅当测点数值发生变化时才产生记录，'
        '因此时间间隔不固定。'
    )

    doc.add_heading('1.2 设备选择', level=2)
    doc.add_paragraph('从矿井全部 8 台设备中选择两台进行阶段一分析：')
    doc.add_paragraph('采煤机（CMJ）—— 核心采掘设备，61 个状态/监测测点', style='List Bullet')
    doc.add_paragraph('转载机（ZZJ）—— 输送系统关键设备，25 个测点', style='List Bullet')

    doc.add_heading('1.3 数据概览对比', level=2)
    add_table(doc,
        ['指标', '采煤机（CMJ）', '转载机（ZZJ）'],
        [
            ['宽表行数',     '39,184',  '87,337'],
            ['宽表列数',     '37',      '15'],
            ['时间跨度',     '27 天（04-01 ~ 04-28）', '61 天（04-01 ~ 05-31）'],
            ['缺失率',       '31 列 < 1%', '11 列 < 1%'],
        ],
        col_widths=[5, 5.5, 5.5],
    )
    spacer(doc)

    # ═══ 2. 数据预处理 ═══
    doc.add_heading('2. 数据预处理', level=1)

    doc.add_heading('2.1 On-change 存储机制', level=2)
    doc.add_paragraph(
        'On-change 存储的特点：数值变化时才记录，隐含' + LQ + '保持' + RQ + '行为——'
        '从记录时间点到下一次变化，数值保持不变。这使得原始数据为长表格式'
        '（point_name / time / value），且时间间隔不均匀。'
    )

    doc.add_heading('2.2 重采样流程', level=2)
    steps = [
        '分块读取 CSV（iter_chunks 策略，避免大文件内存溢出）',
        '剔除无效测点：班次统计、干预状态、冷却水压力等非监测参数',
        '按测点分组，每个测点构成不完整时间序列',
        '前向填充（ffill）：补全测点未记录的时段',
        '重采样到 1 分钟网格：取每个 1 分钟区间最后一个值（last()）',
        '再次前向填充：补全 1 分钟网格中仍为空的值',
        '输出完整等间隔宽表（Parquet 格式）',
    ]
    for s in steps:
        p = doc.add_paragraph(s, style='List Number')
        p.paragraph_format.space_after = Pt(2)

    spacer(doc)
    doc.add_paragraph(
        '选用 last() 而非 mean()：on-change 的语义是' + LQ + '值在该时刻变为新值并保持到下一次变化' + RQ + '，'
        'last() 保留最近有效状态，而 mean() 会抹掉瞬间变化信息。'
    )
    doc.add_paragraph(
        '1 分钟网格的选择依据：煤矿设备物理过程的时间常数在秒到分钟级别，'
        '1 分钟足够捕捉工况切换和电流趋势，同时将数据量压缩到可管理规模。'
    )

    # ═══ 3. 数据空洞检测 ═══
    doc.add_heading('3. 数据空洞检测', level=1)
    doc.add_paragraph(
        'On-change 数据的隐含假设：无变化 = 数值保持不变。但如果传感器断线、'
        '通讯中断或电源故障，数值也会"保持不变"——这形成了数据空洞。'
    )
    doc.add_heading('3.1 检测方法', level=2)
    doc.add_paragraph(
        '在重采样到 1 分钟网格的宽表上，对每个监测参数扫描连续相同值的'
        '游程（run-length encoding）。游程长度 ≥ 120 分钟的区间标记为空洞。'
        '120 分钟的阈值基于经验：正常工况切换和参数漂移可在 2 小时内产生数值变化，'
        '连续 2 小时不变通常意味着数据质量问题。'
    )
    doc.add_paragraph(
        '空洞检测输出：`gap_report.csv`（空洞明细表）和 `gap_detection.png`'
        '（热力图，绿色=数据正常，红色=空洞，颜色深度与空洞时长成正比）。'
    )
    doc.add_heading('3.2 计算结果', level=2)
    doc.add_paragraph(
        '空洞检测结果标注在 gap_detection.png 热力图中，并列出了各参数的空洞时段。'
        '典型空洞模式：深夜到凌晨时段部分参数无变化，可能与交接班或设备待机相关。'
    )

    # ═══ 4. 参数层级图谱与分部位工况 ═══
    doc.add_heading('4. 参数层级图谱与分部位工况', level=1)
    doc.add_heading('4.1 命名规则解析', level=2)
    doc.add_paragraph(
        '采煤机测点遵循标准命名规则：`设备_部位_组件_传感器_指标`。'
        '例如 `采煤机_截割部位_右滚筒_电机_电流` 解析为：'
    )
    add_table(doc,
        ['层级', '值'],
        [['设备（device）', '采煤机'], ['部位（part）', '截割部位'],
         ['组件（component）', '右滚筒'], ['传感器（sensor）', '电机'],
         ['指标（metric）', '电流']],
        col_widths=[4, 8],
    )
    spacer(doc)
    doc.add_paragraph(
        '生成的 `param_hierarchy.csv` 按部位分组展示完整层级关系，'
        '`param_hierarchy.png` 为树状可视化。'
    )

    doc.add_heading('4.2 分部位工况体系', level=2)
    doc.add_paragraph(
        '废弃原有的 L1→L2→L3 三层层次化方案，改为分部位独立工况判定。'
        '每个部位根据自身传感器独立判定，再汇总推导设备级工况。'
        '分部位的优势：截割部"割煤中"时牵引部可能"待机"（斜切进刀场景），'
        '这种局部态在整机工况中会被淹没。'
    )
    parts_table = [
        ['截割部_工况',
         '割煤低位 / 割煤中位 / 割煤高位 / 调架中 / 待机-高位 / 待机 / 停机',
         '任一滚筒运行 + 速度 > 0.5 m/min → 按 max(左右高度) 分：<3m→割煤低位，3~5m→割煤中位，≥5m→割煤高位；'
         '运行 + 速度 ≤ 0.5 → 调架中；滚筒停转 → 左高度<4.5m→待机，≥4.5m→待机-高位'],
        ['牵引部_工况', '空载牵引 / 重载牵引 / 待机 / 停机',
         '牵引电机运行 + 速度 > 0 → 牵引，再按左右电机电流之和 < 100A → 空载牵引，≥ 100A → 重载牵引；运行 + 速度为 0 → 待机'],
        ['油泵_工况', '轻载 / 重载 / 停机',
         '油泵电机运行→ 按平均油压 < 1.0 MPa → 轻载，≥ 1.0 MPa → 重载'],
        ['破碎机_工况', '空载运行 / 带载运行 / 停机',
         '破碎机电机运行→ 按电机电流 < 50A → 空载运行，≥ 50A → 带载运行'],
        ['设备_工况', '割煤中 / 正常运行 / 空载牵引 / 待机 / 停机',
         '从 4 个部位列推导的便捷视图'],
    ]
    add_table(doc,
        ['部门', '状态取值', '判定逻辑'],
        parts_table,
        col_widths=[2.5, 3.5, 6],
    )
    spacer(doc)

    # ═══ 5. 工况划分 ═══
    doc.add_heading('5. 工况划分（分部位）', level=1)

    doc.add_heading('5.1 划分策略：分部位独立判定', level=2)
    doc.add_paragraph(
        '工况划分采用分部位独立判定策略，废弃原有的层次化 L1→L2→L3 方案。'
        '每个部位（截割部/牵引部/油泵/破碎机）依据自身传感器独立标注多状态工况，'
        '再汇总为设备级工况（设备_工况）。'
        '分部位的优势：部位间解耦——截割部"割煤中"时牵引部可"待机"，'
        '这种组合态在传统层次方案中无法表达。'
    )

    doc.add_heading('5.2 截割部_工况：割煤低位 / 割煤中位 / 割煤高位 / 调架中 / 待机-高位 / 待机 / 停机', level=2)
    doc.add_paragraph(
        '判定规则：左右滚筒任一电机运行 + 采煤机速度 > 0.5 m/min → 割煤中，再按滚筒最大高度细分：'
        '< 3m → 割煤低位，3~5m → 割煤中位，≥ 5m → 割煤高位。'
        '运行 + 速度 ≤ 0.5 → 调架中。滚筒停转时，按左滚筒高度区分：'
        '< 4.5m → 待机，≥ 4.5m → 待机-高位。速度阈值 0.5 m/min 基于领域知识——'
        '正常割煤牵引速度 > 2 m/min，调架时机器极慢蠕动或静止。'
    )
    doc.add_paragraph(
        '优化：聚类验证显示左滚筒高度（30.6%）和摇臂角度（20.9%）对截割部'
        '样本分离贡献远超电流与速度——KMeans 的第一分离轴是高度而非电流。'
        '原"割煤中"态在高度空间呈现连续分布（0.05m~6.3m），与 KMeans 聚类的自然边界不对齐。'
        '通过引入高度阈值（3m/5m）将割煤中拆分为低位/中位/高位三个子态，'
        '同时从聚类特征中移除高度特征，使规则维度与聚类特征对齐。'
        '优化后 ARI 从 0.17 提升至 0.59。'
        '当前方案共七态：割煤低位/割煤中位/割煤高位/调架中/待机-高位/待机/停机。'
    )

    doc.add_heading('5.3 牵引部_工况：空载牵引 / 重载牵引 / 待机 / 停机', level=2)
    doc.add_paragraph(
        '判定规则：牵引电机运行 + 速度 > 0 → 牵引；运行 + 速度为 0 → 待机。'
        '牵引部独立于截割部：即使滚筒未运行，只要机器在行走即标记为牵引。'
    )
    doc.add_paragraph(
        '优化：在"牵引"态内，基于左右电机电流之和阈值（< 100A = 空载牵引，'
        '≥ 100A = 重载牵引）细分。原 ARI = 0.57 已较高，电流细分进一步提升'
        '对牵引负载变化的敏感度。'
    )

    doc.add_heading('5.4 油泵_工况：轻载 / 重载 / 停机', level=2)
    doc.add_paragraph(
        '判定规则：油泵电机运行 → 按平均油压是否 ≥ 1.0 MPa 拆分为轻载/重载。'
        '原二值"运行/停机"忽略了油泵在不同负载下油压的差异——'
        '油泵电流在运行态内存在明显的轻载/重载子模式，通过油压阈值分解决此问题。'
    )

    doc.add_heading('5.5 破碎机_工况：空载运行 / 带载运行 / 停机', level=2)
    doc.add_paragraph(
        '判定规则：破碎机电机运行 → 按电机电流 < 50A = 空载运行，≥ 50A = 带载运行。'
        '原二值"运行/停机" ARI = 0.10 说明仅靠运行状态无法反映破碎机的真实工作模式——'
        '电流在运行态内有多簇分布。电流阈值细分后预期 ARI 有显著提升。'
    )

    doc.add_heading('5.6 设备_工况（衍生）：割煤中 / 正常运行 / 空载牵引 / 待机 / 停机', level=2)
    doc.add_paragraph(
        '从 4 个部位列推导的便捷视图。部位子态在设备级按优先级归并——'
        '割煤中最优先，其次正常运行，空载牵引，待机，停机。'
    )
    add_table(doc,
        ['设备工况', '推导条件'],
        [
            ['停机',     '4 个部位全部停机'],
            ['待机',     '非停机且非空载牵引且非截割部运行'],
            ['空载牵引', '牵引部=牵引中且截割部=待机/停机'],
            ['正常运行',   '截割部=调架中'],
            ['割煤中',   '截割部=割煤中'],
        ],
        col_widths=[3, 6],
    )
    spacer(doc)

    doc.add_heading('5.7 转载机工况', level=2)
    doc.add_paragraph('转载机逻辑较简单：运行状态 + 电流阈值。')
    add_table(doc,
        ['工况', '判定条件', '样本数', '占比'],
        [
            ['带载运行', '运行 + 电流 ≥ 50 A',  '51,717', '59.2%'],
            ['停机',     '运行状态 = 0',        '34,682', '39.7%'],
            ['空载运行', '运行 + 电流 < 50 A',  '938',    '1.1%'],
        ],
        col_widths=[2.5, 4, 3, 2],
    )
    spacer(doc)

    # ═══ 6. 分工况统计分析 ═══
    doc.add_heading('6. 分工况统计分析', level=1)

    doc.add_heading('6.1 为什么分工况统计', level=2)
    doc.add_paragraph(
        '同一监测参数在不同工况下表现截然不同。例如截割电流：停机时 0 A，'
        '割煤时均值 92.97 A，调架时均值 189.67 A（被极端值拉高）。'
        '如果不分工况直接计算全局均值和标准差，工况间巨大差异会淹没工况内微小异常——'
        '分工况是后续异常检测的前提。'
    )

    doc.add_heading('6.2 采煤机关键统计发现', level=2)

    p = doc.add_paragraph()
    run = p.add_run('发现 1：调架电流均值（189.67 A）高于割煤（92.97 A）')
    run.font.bold = True
    doc.add_paragraph(
        '原因：调架包含大量' + LQ + '滚筒埋煤 + 速度为零' + RQ + '的情形，滚筒对着煤壁转动但机器不前进，'
        '导致电流骤升（max 1,132 A）。而割煤时正常截割，电流平稳（std 38.49 vs 调架 330.01）。'
        '结论：调架工况用中位数（64.0 A）而非均值作为参考更稳健。',
        style='List Bullet',
    )
    p = doc.add_paragraph()
    run = p.add_run('发现 2：牵引电机左右对称性好')
    run.font.bold = True
    doc.add_paragraph(
        '割煤工况下，右牵引电流均值 124.63 A，左 124.91 A（差异 < 0.3%）；'
        '右温度均值 64.15°C，左 63.72°C（差异 < 1%）。两侧负荷均匀，传动链健康。',
        style='List Bullet',
    )

    doc.add_heading('6.3 转载机统计发现', level=2)
    doc.add_paragraph(
        '带载运行占 59.2%，与下游采煤生产直接相关。'
        '停机占 39.7%（含采煤机停机时期，合理）。'
        '空载运行仅 1.1%，说明启动和轻载时间占比极低。',
        style='List Bullet',
    )

    # ═══ 7. 可视化结果 ═══
    doc.add_heading('7. 可视化结果', level=1)

    doc.add_heading('7.1 原始图表', level=2)
    doc.add_paragraph('以下图表在反馈前已生成，构成阶段一核心可视化集：')

    doc.add_heading('7.1.1 采煤机图表', level=3)
    cmj_charts = [
        ('cmj_device_timeline', '设备级工况（设备_工况）时间线：割煤中/正常运行/空载牵引/待机/停机。'),
        ('cmj_device_pie', '设备级工况占比饼图。'),
        ('cmj_cutting_timeline', '截割部_工况时间线：割煤低位/割煤中位/割煤高位/调架中/待机-高位/待机/停机。'),
        ('cmj_traction_timeline', '牵引部_工况时间线：空载牵引/重载牵引/待机/停机。'),
        ('cmj_pump_timeline', '油泵_工况时间线：轻载/重载/停机。'),
        ('cmj_crusher_timeline', '破碎机_工况时间线：空载运行/带载运行/停机。'),
        ('cmj_cut_current',
         '截割部——割煤中电流均值约 150 A，调架中电流分布极散（均值约 190 A），'
         '说明调架时负载差异极大。割煤中按高度拆分为低位/中位/高位三子态，'
         '但电流散布大，各子态间重叠度高。'),
        ('cmj_trac_current',
         '牵引部——空载牵引/重载牵引电流分布集中，待机时电流接近 0，左右牵引对称性好。'),
        ('cmj_pump_current',
         '油泵——轻载 vs 重载 vs 停机油泵电流对比，电流阈值细分负载状态。'),
        ('cmj_crusher_current',
         '破碎机——空载运行 vs 带载运行 vs 停机破碎机电流对比。'),
        ('cmj_corr', '左右截割电流高度正相关；截割电流与牵引电流中等正相关；电流与温度中等正相关但存在热惯性滞后。'),
    ]
    for key, desc in cmj_charts:
        title, fname = CHART_MAP[key]
        doc.add_heading(title, level=3)
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
        doc.add_paragraph(desc)

    doc.add_heading('7.1.2 转载机图表', level=3)
    zzj_charts = [
        ('zzj_timeline', '带载运行 59.2%，停机 39.7%，空载运行 1.1%。'),
        ('zzj_pie', '带载运行占比与下游生产直接挂钩。'),
        ('zzj_current', '带载/停机/空载三个工况的电流分布分离清晰。'),
        ('zzj_corr', '电机电流与转速正相关，IGBT 温度随负载上升。'),
    ]
    for key, desc in zzj_charts:
        title, fname = CHART_MAP[key]
        doc.add_heading(title, level=3)
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
        doc.add_paragraph(desc)

    doc.add_heading('7.2 新增分析图表（基于反馈）', level=2)
    doc.add_paragraph('以下图表基于 7/8 反馈建议新增，深化了数据质量、参数结构和跨工况分析。')

    new_charts = [
        ('gap_detection',
         '数据空洞热力图。横轴=监测参数，纵轴=时间，绿色=数据正常，红色=空洞区间'
         '（颜色深度与持续时长成正比）。连续相同值 ≥ 120 min 标记为空洞。'),
        ('param_hierarchy',
         '参数层级图谱。按 设备_部位_组件_传感器_指标 命名规则解析，'
         '分部位（截割部/牵引部/油泵/破碎机）展示完整层级树。'),
        ('cmj_profile',
         '跨工况参数 Profile 对比。z-score 归一化后展示同一参数在设备_工况'
         '（割煤中/正常运行/空载牵引/待机/停机）下的均值±标准差，'
         '支持跨参数比较各工况的特征差异。'),
        ('cluster_vs_截割部',
         '截割部——聚类 vs 规则工况对比。左：混淆矩阵（归一化），'
         '右：PCA 降维散点图。Adjusted Rand Index 量化规则标签与 KMeans 聚类结果的一致性。'),
        ('cluster_vs_牵引部',
         '牵引部——聚类 vs 规则工况对比。'),
        ('cluster_vs_油泵',
         '油泵——聚类 vs 规则工况对比。'),
        ('cluster_vs_破碎机',
         '破碎机——聚类 vs 规则工况对比。'),
    ]
    for key, desc in new_charts:
        title, fname = CHART_MAP[key]
        doc.add_heading(title, level=3)
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
        doc.add_paragraph(desc)

    doc.add_heading('7.3 参数滞后互相关图', level=2)
    doc.add_paragraph(
        '滞后互相关（lagged cross-correlation）分析关键参数对间的时序耦合关系，'
        '使用 Spearman 秩相关，滞后范围 [-30, +30] 分钟。'
    )
    lagged_pairs = [
        ('截割电流 vs 截割温度', 'cmj_lagged_corr_电流_温度.png'),
        ('牵引电流 vs 牵引速度', 'cmj_lagged_corr_牵引电流_速度.png'),
        ('俯仰角 vs 位置架号', 'cmj_lagged_corr_俯仰角_架号.png'),
    ]
    for label, fname in lagged_pairs:
        doc.add_heading(label, level=3)
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
        doc.add_paragraph(
            '每条曲线对应一个设备_工况。x 轴=滞后时间（负值=参数 A 领先），'
            'y 轴=Spearman 相关系数。峰值的偏移量反映两参数间的响应延迟。'
        )

    # ═══ 8. 工况转换分析 ═══
    doc.add_heading('8. 工况转换分析（分部位）', level=1)

    doc.add_heading('8.1 分析方法', level=2)
    doc.add_paragraph(
        '工况转换分析对每个部位独立检测切换事件。'
        '方法：在工况列上计算 shift(1) != shift(-1) 定位切换帧，'
        '记录切换前后工况、持续帧数、以及窗口内各参数的均值。'
        '分析覆盖 5 个工况序列：设备_工况、截割部_工况、牵引部_工况、'
        '油泵_工况、破碎机_工况。'
    )

    doc.add_heading('8.2 主要发现', level=2)

    p = doc.add_paragraph()
    run = p.add_run('截割部切换频率')
    run.font.bold = True
    doc.add_paragraph(
        '截割部工况切换集中在割煤各子态↔调架中、割煤各子态↔待机和待机↔待机-高位。'
        '割煤（低位/中位/高位）→调架中切换发生在工作面两端转向时，典型特征：速度降低、'
        '截割电流下降、摇臂角度调整。割煤→待机切换发生在长时停机前，'
        '各参数先后归零（截割电流先降，温度后降——热惯性）。'
        '待机↔待机-高位切换反映滚筒高度调整——低位待机（<4.5m）和高位待机（≥4.5m）'
        '在散点图上形成分离的两个聚簇。',
        style='List Bullet',
    )
    p = doc.add_paragraph()
    run = p.add_run('牵引部切换频率')
    run.font.bold = True
    doc.add_paragraph(
        '牵引中↔待机切换对应采煤机启停操作。牵引中→待机典型速度从工作值'
        '（2-6 m/min）归零；待机→牵引中从 0 爬升到目标速度，伴随牵引电流上升。',
        style='List Bullet',
    )
    p = doc.add_paragraph()
    run = p.add_run('油泵 / 破碎机切换')
    run.font.bold = True
    doc.add_paragraph(
        '切换窗口中，电流在 1-2 分钟内完成阶跃（响应迅速），'
        '温度在 5-15 分钟内逐渐变化（热惯性主导）。'
        '这种响应时差是多维异常检测中重要的特征维度。',
        style='List Bullet',
    )

    doc.add_paragraph(
        '转换分析输出（分部位）：'
    )
    doc.add_paragraph(
        '• transition_设备_工况_stats.csv —— 设备级工况切换统计',
        style='List Bullet',
    )
    doc.add_paragraph(
        '• transition_截割部_工况_stats.csv —— 截割部切换统计',
        style='List Bullet',
    )
    doc.add_paragraph(
        '• transition_牵引部_工况_stats.csv —— 牵引部切换统计',
        style='List Bullet',
    )
    doc.add_paragraph(
        '• transition_油泵_工况_stats.csv —— 油泵切换统计',
        style='List Bullet',
    )
    doc.add_paragraph(
        '• transition_破碎机_工况_stats.csv —— 破碎机切换统计',
        style='List Bullet',
    )
    doc.add_paragraph(
        '切换时序图（transition_电流.png、transition_速度.png、transition_角度.png）'
        '标注切换点前后各 10 分钟窗口。'
    )

    # ═══ 9. 聚类验证规则工况（分部位） ═══
    doc.add_heading('9. 聚类验证规则工况（分部位）', level=1)

    doc.add_heading('9.1 方法', level=2)
    doc.add_paragraph(
        '对每个部位独立做聚类验证。以该部位的关键监测参数为特征'
        '（截割部→截割电流/滚筒高度/摇臂角度；牵引部→牵引电流/速度；'
        '油泵→油泵电流；破碎机→破碎机电流），StandardScaler 标准化后使用 KMeans 聚类。'
        '聚类数 k 取该部位规则标签类别数。'
        '通过 Adjusted Rand Index（ARI）量化聚类结果与规则标签的一致性。'
    )

    doc.add_heading('9.2 结果分析', level=2)
    doc.add_paragraph(
        'ARI 取值范围 [-1, 1]，1 表示完全一致，0 表示随机分配。'
        '实际 ARI 值高 → 规则划分与数据自然聚类高度吻合，'
        '规则阈值的选取合理。'
    )
    doc.add_paragraph(
        '分部位结果解读：'
    )
    doc.add_paragraph(
        '截割部——速度阈值 0.5 m/min 对"割煤中"与"调架中"的划分效果最佳。'
        '不一致样本集中在速度 0.3-0.7 m/min 的模糊区域。',
        style='List Bullet',
    )
    doc.add_paragraph(
        '牵引部——"牵引中"与"待机"边界以速度为 0 为界，聚类一致性好。',
        style='List Bullet',
    )
    doc.add_paragraph(
        '油泵 / 破碎机——二值工况聚类分离度高，ARI 趋于 1。',
        style='List Bullet',
    )
    doc.add_paragraph(
        '输出：每个部位一张 `cluster_vs_截割部.png` / `cluster_vs_牵引部.png` / '
        '`cluster_vs_油泵.png` / `cluster_vs_破碎机.png`（混淆矩阵 + PCA 散点图）'
        '及对应的 `cluster_disagreement_截割部.csv` 分歧样本明细。'
    )

    # ═══ 10. 参数关联深入与跨工况对比 ═══
    doc.add_heading('10. 参数关联深入与跨工况对比', level=1)

    doc.add_heading('10.1 滞后互相关分析', level=2)
    doc.add_paragraph(
        '传统 Spearman 相关分析（热力图）只考虑同期相关性，'
        '但煤矿设备参数间存在明显的响应延迟。滞后互相关通过滑动时间偏移'
        '来量化这种延迟耦合：'
    )
    doc.add_paragraph(
        '截割电流 vs 截割温度：零滞后处相关系数约 0.6-0.7，'
        '但温度电流的峰值相关出现在滞后 +5~+10 分钟——'
        '温度上升滞后于电流变化，这符合热传导物理过程。',
        style='List Bullet',
    )
    doc.add_paragraph(
        '牵引电流 vs 牵引速度：零滞后处相关系数约 0.7-0.8，'
        '且不同工况下曲线形态差异显著——割煤时电流-速度同步性高，'
        '调架时解耦。',
        style='List Bullet',
    )
    doc.add_paragraph(
        '俯仰角 vs 位置架号：低相关（|r| < 0.3），'
        '说明俯仰角受地形影响大于受位置影响，'
        '或当前量测精度不足以揭示弱相关。',
        style='List Bullet',
    )

    doc.add_heading('10.2 跨工况参数 Profile 对比', level=2)
    doc.add_paragraph(
        '将采煤机所有关键参数按设备_工况分组，计算每组均值±标准差，'
        'z-score 归一化后绘制 grouped bar chart。'
        '该图直观展示了各工况下参数的"特征指纹"：'
    )
    doc.add_paragraph(
        '割煤中：截割电流、牵引电流、速度同步升高——典型生产态', style='List Bullet')
    doc.add_paragraph(
        '调架中：截割电流极端高（滚筒埋煤）但速度极低——非生产态', style='List Bullet')
    doc.add_paragraph(
        '空载牵引：速度中高但电流低——仅有行走动作', style='List Bullet')
    doc.add_paragraph(
        '待机：除温度和振动外几乎全低——通电待命', style='List Bullet')
    doc.add_paragraph(
        '这种"工况指纹"为后续异常检测提供了基线模板：实时参数偏离'
        '当前工况的指纹模式即可触发告警。'
    )

    # ═══ 11. 关键发现 ═══
    doc.add_heading('11. 关键发现', level=1)

    doc.add_heading('11.1 状态量-监测参数关系', level=2)
    doc.add_paragraph('回答任务书核心问题：哪些状态量（组合）影响哪些数值监测参数？')
    add_table(doc,
        ['状态量组合', '影响参数', '影响模式'],
        [
            ['滚筒运行状态 + 采煤机速度', '截割电流',
             '割煤时电流激增 3-5 倍，速度与电流非单调'],
            ['牵引方向 + 位置架号', '牵引电流',
             '上行/下行电流差异小，电气对称性好'],
            ['记忆割煤状态', '电流波动模式',
             '自动化切割时电流波动更规则稳定'],
            ['转载机运行状态', '电机电流 + 转速',
             '运行-电流强相关；堵料时高转速 + 低电流'],
            ['截割状态', '电机温度',
             '温度上升滞后截割开始约 5-15 min（热惯性）'],
        ],
        col_widths=[4.5, 3.5, 5.5],
    )
    spacer(doc)

    doc.add_heading('11.2 工况分布意义', level=2)
    doc.add_paragraph(
        '采煤机运行率 71.2%，停机 28.8%（含交接班、检修、故障停机）。'
        '割煤占运行期 64.1%，待机占 20.6%——待机比例值得关注，'
        '可能是生产组织优化的潜在方向。'
    )
    doc.add_paragraph(
        '转载机带载率 59.2%，停机 39.7%。停机比例高部分原因是采煤机停机期间转载机必然停机，'
        '二者联动在阶段三多设备分析中需进一步量化。'
    )

    # ═══ 12. 阶段二预备 ═══
    doc.add_heading('12. 阶段二预备', level=1)

    doc.add_paragraph('阶段一的输出直接服务于阶段二的异常检测建模：')
    for item in [
        '工况标记宽表 → 阶段二在其上做滑动窗口特征提取',
        '分工况统计量（mean / std / median）→ 作为 3σ / IQR 基线的基准值',
        '参数间 Spearman 相关关系 → 指导马氏距离的维度选择',
        '工况时间线 → 辅助验证异常事件是否发生在工况切换时刻',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('阶段二计划方法', level=2)
    add_table(doc,
        ['方法', '类型', '可解释性', '优先级'],
        [
            ['分工况 3σ / IQR',     '统计基线',  '高',   '1（首先实现）'],
            ['滑动窗口 + 马氏距离', '多维统计',  '中',   '2'],
            ['Isolation Forest',    '机器学习',  '中',   '3（补充）'],
            ['LSTM-AE / VAE',       '深度学习',  '低',   '4（可选）'],
        ],
        col_widths=[4.5, 3, 2.5, 3],
    )
    spacer(doc)

    doc.add_paragraph('重点异常指标设计：')
    for a in [
        '截割电流相对工况基线的偏差（z-score）',
        '牵引电机温升率异常（导数的标准差监控）',
        '截割电流-牵引电流联合异常（马氏距离）',
        '转载机 IGBT 温升 vs 电流不匹配（堵料检测）',
    ]:
        doc.add_paragraph(a, style='List Bullet')

    # ═══ 附录 ═══
    doc.add_page_break()
    doc.add_heading('附录：图表清单与说明', level=1)
    add_table(doc,
        ['图表名称', '类型', '核心信息'],
        [
            # 原始图表（10+ 张）
            ['设备级工况时间线',             '甘特图',   '设备_工况（割煤中/正常运行/空载牵引/待机/停机）切换'],
            ['设备级工况占比',               '饼图',     '设备_工况占比分布'],
            ['截割部工况时间线',             '甘特图',   '截割部_工况（割煤低位/割煤中位/割煤高位/调架中/待机-高位/待机/停机）切换'],
            ['牵引部工况时间线',             '甘特图',   '牵引部_工况（空载牵引/重载牵引/待机/停机）切换'],
            ['油泵工况时间线',               '甘特图',   '油泵_工况（轻载/重载/停机）切换'],
            ['破碎机工况时间线',             '甘特图',   '破碎机_工况（空载运行/带载运行/停机）切换'],
            ['截割部截割电流箱线图',         '箱线图',   '分工况（割煤低位/割煤中位/割煤高位/调架中/待机-高位/待机/停机）截割电流分布'],
            ['牵引部牵引电流箱线图',         '箱线图',   '分工况（空载牵引/重载牵引/待机/停机）牵引电流分布'],
            ['油泵电流箱线图',               '箱线图',   '轻载 vs 重载 vs 停机油泵电流对比'],
            ['破碎机电流箱线图',             '箱线图',   '空载运行 vs 带载运行 vs 停机破碎机电流对比'],
            ['CMJ Spearman 热力图',          '热力图',   '参数间单调相关性'],
            ['ZZJ 工况时间线',               '甘特图',   '带载/停机/空载切换'],
            ['ZZJ 工况占比',                 '饼图',     '带载 59.2% 为主工况'],
            ['ZZJ 分工况电流箱线图',         '箱线图',   '三工况电流分布分离清晰'],
            ['ZZJ Spearman 热力图',          '热力图',   '电流-温度-转速相关性'],
            # 新增分析图表（基于 7/8 反馈）
            ['数据空洞热力图',               '热力图',   '连续相同值 ≥ 120 min 标记为空洞'],
            ['参数层级图谱',                 '树状图',   '设备→部位→组件→传感器→指标层级'],
            ['跨工况参数 Profile 对比',      '柱状图',   'z-score 归一化跨工况特征指纹（设备_工况）'],
            ['截割部聚类 vs 规则工况对比',   '混合图',   '混淆矩阵 + PCA 散点，分部位 ARI'],
            ['牵引部聚类 vs 规则工况对比',   '混合图',   '混淆矩阵 + PCA 散点，分部位 ARI'],
            ['油泵聚类 vs 规则工况对比',     '混合图',   '混淆矩阵 + PCA 散点，分部位 ARI'],
            ['破碎机聚类 vs 规则工况对比',   '混合图',   '混淆矩阵 + PCA 散点，分部位 ARI'],
            ['截割电流滞后互相关',           '折线图',   '电流 vs 温度，滞后 ±30 min Spearman'],
            ['牵引电流滞后互相关',           '折线图',   '牵引电流 vs 速度，分工况曲线'],
            ['俯仰角滞后互相关',             '折线图',   '俯仰角 vs 位置架号，弱相关验证'],
            ['切换时序图（电流）',           '折线图',   '切换窗口内电流变化'],
            ['切换时序图（速度）',           '折线图',   '切换窗口内速度变化'],
            ['切换时序图（角度）',           '折线图',   '切换窗口内角度变化'],
        ],
        col_widths=[5, 2, 6.5],
    )

    # ── 保存 ──
    os.makedirs(PHASE1, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f'[OK] 报告已生成: {OUTPUT_FILE}')
    print(f'    文件大小: {os.path.getsize(OUTPUT_FILE) / 1024:.0f} KB')


# ═══════════════════════════════════════════
# 第二周报告：深层分析（不含基础预处理/可视化）
# ═══════════════════════════════════════════

W2_OUTPUT_FILE = os.path.join(PHASE1, '第二周_深度分析报告.docx')


def build_week2_report():
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Microsoft YaHei'
        hs.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        if level == 1:
            hs.font.size = Pt(18)
            hs.paragraph_format.space_before = Pt(18)
            hs.paragraph_format.space_after = Pt(10)
        elif level == 2:
            hs.font.size = Pt(14)
            hs.paragraph_format.space_before = Pt(14)
            hs.paragraph_format.space_after = Pt(6)
        else:
            hs.font.size = Pt(12)
            hs.paragraph_format.space_before = Pt(10)
            hs.paragraph_format.space_after = Pt(4)

    add_page_number(doc)

    # ═══ 封面 ═══
    for _ in range(5):
        spacer(doc, 14)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run('第二周：深度分析报告')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_line.add_run('━' * 32)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    spacer(doc, 4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('采煤机（CMJ）+ 转载机（ZZJ）工况深层分析')
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run('非参数检验 · 段持续统计 · 聚类验证 · 切换时域 · 转载机深度分析')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    spacer(doc, 16)

    info_lines = [
        ('数据来源', '大海则煤矿  2024-04-01 ~ 2024-04-28（CMJ）/ 2024-05-31（ZZJ）'),
        ('分析设备', '采煤机（CMJ）+ 转载机（ZZJ）'),
        ('工况体系', 'CMJ 分部位（截割部/牵引部/油泵/破碎机）独立判定，ZZJ 单一工况列'),
        ('报告日期', '2026-07-14'),
    ]
    for label, value in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f'{label}：')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run = p.add_run(value)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_page_break()

    # ══ 目录 ══
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 当前工况分类规则与参数层级图谱',
        '2. Kruskal-Wallis 非参数显著性检验',
        '3. 状态段持续时间与异常段检测',
        '4. 分部位聚类验证规则工况',
        '5. 工况转换时域特征分析与聚合剖面',
        '6. 分部位滞后互相关分析',
        '7. 转载机深度分析',
        '8. 关键发现与阶段二预备',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(14), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
        run = p.add_run(item)
        run.font.size = Pt(12)
        run.font.bold = True

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 1. 当前工况分类规则与参数层级图谱
    # ═══════════════════════════════════════════════════════════
    doc.add_heading('1. 当前工况分类规则与参数层级图谱', level=1)

    doc.add_heading('1.1 四部位独立工况体系', level=2)
    doc.add_paragraph(
        '本项目采用分部位独立工况判定的方法：将采煤机（CMJ）分为四个机'
        '械部件——截割部、牵引部、油泵、破碎机——各自基于关键监测参数独立'
        '计算工况，最终汇总为设备级综合工况。这一设计取代了传统的单一 L1/L2/L3 '
        '层级体系，使工况精细度大幅提升。'
    )

    doc.add_heading('1.2 各部位工况划分规则', level=2)

    rules_data = [
        ['部件', '工况状态', '判定依据', '主要参数'],
        ['截割部', '割煤低位', '滚筒运行 + 速度>0.5m/min + max(高度)<3m',
         '滚筒高度、滚筒转速、牵引速度、截割电流'],
        ['截割部', '割煤中位', '滚筒运行 + 速度>0.5m/min + 3m≤max(高度)<5m',
         '滚筒高度、滚筒转速、牵引速度、截割电流'],
        ['截割部', '割煤高位', '滚筒运行 + 速度>0.5m/min + max(高度)≥5m',
         '滚筒高度、滚筒转速、牵引速度、截割电流'],
        ['截割部', '调架中', '滚筒运行 + 速度 ≤ 0.5 m/min',
         '滚筒转速、牵引速度'],
        ['截割部', '待机', '滚筒运行 + 速度为 0 + 左高度<4.5m',
         '滚筒高度、滚筒转速'],
        ['截割部', '待机-高位', '滚筒运行 + 速度为 0 + 左高度≥4.5m',
         '滚筒高度、滚筒转速'],
        ['截割部', '停机', '滚筒停止（电流 < 阈值）',
         '截割电机电流'],
        ['牵引部', '空载牵引', '牵引电机运行 + 速度 > 0 + 左右电流和 < 100A',
         '牵引速度、左右牵引电流'],
        ['牵引部', '重载牵引', '牵引电机运行 + 速度 > 0 + 左右电流和 ≥ 100A',
         '牵引速度、左右牵引电流'],
        ['牵引部', '待机', '牵引电机运行 + 速度为 0',
         '牵引速度'],
        ['牵引部', '停机', '牵引电机电流 < 阈值',
         '牵引电机电流'],
        ['油泵', '轻载', '油泵电机运行 + 平均油压 < 1.0 MPa',
         '油泵电机电流、油压'],
        ['油泵', '重载', '油泵电机运行 + 平均油压 ≥ 1.0 MPa',
         '油泵电机电流、油压'],
        ['油泵', '停机', '油泵电机电流 ≤ 阈值',
         '油泵电机电流'],
        ['破碎机', '空载运行', '破碎机电机运行 + 电流 < 50A',
         '破碎机电机电流'],
        ['破碎机', '带载运行', '破碎机电机运行 + 电流 ≥ 50A',
         '破碎机电机电流'],
        ['破碎机', '停机', '破碎机电机电流 ≤ 阈值',
         '破碎机电机电流'],
    ]
    tbl = doc.add_table(rows=len(rules_data), cols=4)
    tbl.style = 'Light List Accent 1'
    for i, row_data in enumerate(rules_data):
        for j, val in enumerate(row_data):
            cell = tbl.rows[i].cells[j]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(8) if i > 0 else Pt(8)
                    if i == 0:
                        run.font.bold = True

    add_caption(doc, '表1  各部位工况分类规则定义')

    doc.add_heading('1.3 参数层级图谱', level=2)
    doc.add_paragraph(
        '参数层级图谱展示了从原始传感器数据到工况判定的完整路径：'
        '原始监测参数 → 部位级工况 → 设备级综合工况。'
        '下图以设备-部位-参数三层结构呈现了当前的监测体系。'
    )
    # Try W2_CHART_MAP first, then fallback to CHART_MAP
    hier_key = next((k for k in W2_CHART_MAP if 'hierarchy' in k or 'param' in k), None)
    hier_img = ''
    if hier_key:
        _, fname = W2_CHART_MAP[hier_key]
        hier_img = os.path.join(PHASE1, fname)
    if not os.path.exists(hier_img):
        hier_key2 = next((k for k in CHART_MAP if 'hierarchy' in k or 'param' in k), None)
        if hier_key2:
            _, fname = CHART_MAP[hier_key2]
            hier_img = os.path.join(PHASE1, fname)
    if not os.path.exists(hier_img):
        hier_img = os.path.join(PHASE1, 'param_hierarchy.png')
    if os.path.exists(hier_img):
        doc.add_picture(hier_img, width=Cm(15))
        p = doc.paragraphs[-1]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, '图1  参数层级图谱 — 设备 → 部位 → 监测参数结构')

    doc.add_heading('1.4 设备级综合工况规则', level=2)
    doc.add_paragraph(
        '设备级综合工况（设备_工况）由四个部位工况按照就高不就低原则组合推导：'
        '只要任一部件处于停机状态，设备级即标记为停机；均在运行中则标记为正'
        '常运行。具体组合规则见下表：'
    )
    device_rules = [
        ['设备级工况', '截割部', '牵引部', '油泵', '破碎机'],
        ['正常运行', '割煤低位/割煤中位/割煤高位/调架中', '空载牵引/重载牵引', '轻载/重载', '空载运行/带载运行'],
        ['停机', '停机', '停机', '停机', '停机'],
    ]
    tbl2 = doc.add_table(rows=len(device_rules), cols=5)
    tbl2.style = 'Light List Accent 1'
    for i, row_data in enumerate(device_rules):
        for j, val in enumerate(row_data):
            cell = tbl2.rows[i].cells[j]
            cell.text = val
            for cp in cell.paragraphs:
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cp.runs:
                    run.font.size = Pt(8)
                    if i == 0:
                        run.font.bold = True
    add_caption(doc, '表2  设备级综合工况组合规则')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 2. Kruskal-Wallis 非参数显著性检验
    # ═══════════════════════════════════════════════════════════
    doc.add_heading('2. Kruskal-Wallis 非参数显著性检验', level=1)

    doc.add_heading('2.1 方法原理', level=2)
    doc.add_paragraph(
        'Kruskal-Wallis 检验是单因素方差分析（ANOVA）的非参数替代方法。'
        '与 ANOVA 假设正态分布和方差齐性不同，Kruskal-Wallis 仅要求数据独立且'
        '有序（ordinal），适用于煤矿传感器数据常见的偏态、重尾分布。'
    )
    doc.add_paragraph(
        '基本原理：将所有组的观测值混合后排序（rank），然后比较各组的平均秩次差异。'
        '若各组来自同一分布，组间平均秩次差异应仅由随机波动引起；'
        '若差异过大，则拒绝原假设（所有组分布相同）。'
    )
    doc.add_paragraph(
        'H 统计量的计算公式：\n'
        'H = (12 / (N(N+1))) \xd7 Σ(Rᵢ² / nᵢ) - 3(N+1)\n'
        '其中 N 为总样本数，nᵢ 为第 i 组的样本数，Rᵢ 为第 i 组的秩和。'
        '当组数 ≥ 3 时，H 近似服从 χ²(k-1) 分布。'
    )
    doc.add_paragraph(
        '效应量指标：epsilon²（ε²）= H / (N² - 1) \xd7 (N+1)/(N-1) \xd7 (N-k)/k，'
        '衡量组间差异占总变异比例，0-1 范围，越大表示分组变量解释力越强。'
        '多重比较校正：Benjamini-Hochberg FDR（False Discovery Rate）'
        '控制假阳性率在 5%。'
    )

    doc.add_heading('2.2 结果摘要', level=2)
    doc.add_paragraph(
        '对 4 个部位分别做 Kruskal-Wallis 检验：以部位工况为分组变量，'
        '检验每个监测参数在各工况间是否存在显著差异。'
    )
    doc.add_paragraph(
        '全部 21 个监测参数的 FDR 校正 p 值均 < 0.05，'
        '说明分部位工况划分能够显著区分各监测参数的分布。'
    )

    # Part-by-part KW results
    kw_parts_summary = [
        ('截割部（7 个工况）',
         'H 统计量最高：采煤机速度（33,350.0）、左右截割电流（31,612~32,303）'
         '——速度与电流是最能区分截割部工况的参数。'
         '最低：俯仰角（265.6），俯仰角受地形影响，与截割工况无关。'
         'KW 全部显著验证了以速度和高程阈值为基础的工况划分有效性。'),
        ('牵引部（4 个工况）',
         'H 统计量最高：采煤机速度（33,350.0）、左右牵引电流（30,355~30,587）。'
         '牵引电流在重载牵引/空载牵引/待机/停机间差异显著。'
         '基于左右牵引电流和实施了电流阈值细分（100A），'
         '将牵引拆分为空载牵引/重载牵引。'),
        ('油泵（3 个工况）',
         'H 统计量最高：左右油泵电流（27,608~27,670）。'
         '油泵工况开关特征明显，电流是天然区分器。'
         '基于油压阈值（1.0 MPa）将运行拆分为轻载/重载。'),
        ('破碎机（3 个工况）',
         'H 统计量最高：破碎机电流（5,612.5）。'
         '破碎机电流在空载运行/带载运行/停机间差异显著。'
         '基于电机电流阈值（50A）将运行拆分为空载运行/带载运行。'),
    ]
    for title, desc in kw_parts_summary:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        doc.add_paragraph(desc, style='List Bullet')

    doc.add_heading('2.3 ε²效应量热力图', level=2)
    doc.add_paragraph(
        '以下热力图展示各部位各监测参数的 ε² 效应量（effect size），'
        '使用 Cohen 标准：ε² ≥ 0.14 为大效应，0.06~0.14 为中等效应，'
        '< 0.06 为小效应。'
    )
    doc.add_paragraph(
        '图表解读方法（双面板）\n'
        '主热力图（左）：每个格子显示 z-score 归一化的工况条件中位数（行内归一化），'
        '数值含义是该参数在该工况下相对于全参平均值偏离了多少个标准差。正值（红色）表示高于平均，'
        '负值（蓝色）表示低于平均。格子边标的 * 号表示该参数在 Kruskal-Wallis 检验中经 FDR 校正后'
        '仍显著（p < 0.05）。图中带 * 号的参数才是真正被工况显著影响的参数。\n\n'
        'ε²效应量柱（右）：每行参数右侧的灰色/彩色柱子显示 Friedman 风格的 ε² 效应量'
        '（基于 KW 的 H 统计量计算）。柱子的灰度/颜色按 Cohen 标准分档：'
        '深灰 = 大效应（ε² ≥ 0.14，对应 H 统计量大），浅灰 = 中效应（0.06~0.14），'
        '极浅 = 小或无效应（< 0.06）。柱长直接反映 "该参数能被工况区分到何种程度"——'
        '柱越长，工况对该参数的影响力越大。\n\n'
        '读图顺序：先从右柱看哪些参数的 ε² 最大（这些是工况区分的关键参数），'
        '再从左热力图看这些参数在各工况下的具体偏离方向（正值还是负值），'
        '从而理解工况划分的物理依据。'
    )
    for part in ['截割部', '牵引部', '油泵', '破碎机']:
        key = f'kw_{part}'
        title, fname = W2_CHART_MAP[key]
        doc.add_heading(f'{part}', level=3)
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)

    doc.add_heading('2.4 结果解读与工况有效性', level=2)
    doc.add_paragraph('从 KW 检验结果可以得出以下关键判断：')
    bullets_kw = [
        '截割部：全部 22 个参数的 FDR 校正 p 值均 < 0.05（H 统计量最低的俯仰角也达 265.6），'
         '验证了以速度阈值为基础的工况划分体系整体有效。速度（H=33,350）和左右截割电流'
         '（H=31,612~32,303）是最强区分参数，俯仰角（H=265.6）受地形而非工况影响',
        '牵引部：采煤机速度 H=33,350 与左右牵引电流 H=30,355~30,587 均为极高值，'
         '说明速度与电流对区分空载/重载牵引效果最显著。'
         '基于电流和阈值 100A 将牵引态拆分为空载/重载牵引，有充分的统计支撑',
        '油泵：左右油泵电流（H=27,608~27,669）和左右油箱油压（H=26,300~26,402）'
         '均为极高 H 统计量——停机/轻载/重载间的参数分布差异极为显著，'
         '油压阈值 1.0 MPa 的拆分与 KW 结果一致',
        '破碎机：H 统计量整体低于其他部位（破碎机电流 H=5,612.5，电机温度 H=3,606.4），'
         '但仍全部显著。基于电流阈值 50A 拆分为空载运行/带载运行是合理的第一步细分',
        '整体而言，全部 21 个监测参数经 FDR 校正后 p < 0.05，'
         '分部位工况划分统计上有效。基于 KW 显著参数实施的工况细化优化'
         '（牵引电流/油压/破碎机电流阈值）已在阶段一最终版本中落地',
    ]
    for b in bullets_kw:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 3. 状态段持续时间与异常段检测
    # ═══════════════════════════════════════════════════════════
    doc.add_heading('3. 状态段持续时间与异常段检测', level=1)

    doc.add_heading('3.1 方法原理——RLE 游程编码', level=2)
    doc.add_paragraph(
        'Run-Length Encoding（RLE / 游程编码）：扫描工况列的时间序列，'
        '将连续相同的工况值编码为"（工况，持续长度，起始时间）"三元组。'
        '这种无损压缩编码在信号处理和数据分析中广泛使用——'
        'PNG 图片压缩就基于 RLE 变体。'
    )
    doc.add_paragraph(
        '在煤矿设备数据中的意义：工况持续时间反映了设备的操作节奏和生产效率。'
        '比如割煤中段持续过短可能意味着频繁启停或控制不稳定，'
        '段持续过长则可能是满载长壁面切割。'
    )
    doc.add_paragraph(
        '统计指标：段数（N）、平均持续（Mean）、中位持续（Median）、'
        '最长/最短持续（Max/Min）、总时长（Sum）。'
        '中位相比均值更适合描述偏态分布——某些工况存在极长/极短的离群段。'
    )

    doc.add_heading('3.2 段持续时间统计结果', level=2)

    # 截割部
    doc.add_heading('截割部_工况', level=3)
    add_table(doc,
        ['工况', '段数', '最短(min)', '最长(min)', '平均(min)', '中位(min)', '总时长(min)'],
        [
            ['停机',       '144',   '1', '233',  '19.8',  '3',  '2,851'],
            ['割煤低位',   '560',   '1',  '71',   '3.2',  '2',  '1,792'],
            ['割煤中位', '1,118',   '1', '228',   '5.1',  '3',  '5,702'],
            ['割煤高位', '1,192',   '1', '414',   '8.7',  '5', '10,391'],
            ['待机',     '1,052',   '1', '418',  '13.6',  '2', '14,278'],
            ['待机-高位',  '185',   '1', '107',   '8.2',  '3',  '1,517'],
            ['调架中',   '2,588',   '1', '476',   '1.6',  '1',  '4,170'],
        ],
        col_widths=[2.5, 2, 2, 2.5, 2.5, 2.5, 3],
    )
    spacer(doc)
    doc.add_paragraph(
        '调架中段数极多（2,588）但中位仅 1 分钟——调架是频繁的短时操作，'
        '每次调架通常只持续不到 1 分钟就完成姿态调整。'
        '割煤低位多短段（平均 3.2 min），对应浅截深调节；'
        '割煤高位段最长（最大 414 分钟，≈ 6.9 小时），对应整面长壁切割。',
        style='List Bullet',
    )

    # 牵引部
    doc.add_heading('牵引部_工况', level=3)
    add_table(doc,
        ['工况', '段数', '最短(min)', '最长(min)', '平均(min)', '中位(min)', '总时长(min)'],
        [
            ['停机',   '1,565', '1', '476',  '8.4', '2', '13,140'],
            ['待机',   '1,428', '1', '312',  '5.0', '1',  '7,096'],
            ['空载牵引', '409', '1',  '11',  '1.1', '1',    '447'],
            ['重载牵引', '2,522', '1', '120',  '7.3', '4', '18,501'],
        ],
        col_widths=[2.5, 2, 2, 2.5, 2.5, 2.5, 3],
    )
    spacer(doc)
    doc.add_paragraph(
        '重载牵引段最长 120 分钟——超过 2 小时的连续牵引对应长壁面完整一刀采煤。'
        '空载牵引 409 段但中位仅 1 分钟——多为短距移动调整（如换向、对滚筒位置）。'
        '重载牵引中位 4 分钟 vs 空载中位 1 分钟，负载差异清晰。',
        style='List Bullet',
    )

    # 油泵
    doc.add_heading('油泵_工况', level=3)
    add_table(doc,
        ['工况', '段数', '最短(min)', '最长(min)', '平均(min)', '中位(min)', '总时长(min)'],
        [
            ['停机', '1,233', '1', '313',  '8.6',  '2', '10,594'],
            ['轻载',   '236', '1', '312', '25.4',  '1',  '5,987'],
            ['重载', '1,178', '1', '496', '19.2',  '8', '22,603'],
        ],
        col_widths=[2.5, 2, 2, 2.5, 2.5, 2.5, 3],
    )
    spacer(doc)
    doc.add_paragraph(
        '油泵重载中位 8 分钟 vs 轻载中位 1 分钟——重载段明显更长，'
        '对应油泵在截割/牵引等高负载工况下的持续压力供油。'
        '轻载段 236 段，多出现在待机/调架时的辅助运行。',
        style='List Bullet',
    )
    # 破碎机
    doc.add_heading('破碎机_工况', level=3)
    add_table(doc,
        ['工况', '段数', '最短(min)', '最长(min)', '平均(min)', '中位(min)', '总时长(min)'],
        [
            ['停机',     '858', '1', '584', '25.8', '8.5', '22,106'],
            ['空载运行', '927', '1', '2,667', '18.3', '5.0', '16,966'],
            ['带载运行', '107', '1',   '2',  '1.0', '1.0',    '112'],
        ],
        col_widths=[2.5, 2, 2, 2.5, 2.5, 2.5, 3],
    )
    spacer(doc)
    doc.add_paragraph(
        '破碎机带载运行仅 107 段（总时长 112 分钟），中位仅 1 分钟——'
        '破碎机带载持续时间极短，多为瞬间过载，正常工况下破碎机以空载运行为主。'
        '空载运行最大段 2,667 分钟（≈ 44 小时）为异常值，'
        '可能是数据边界处截断导致或通信中断伪持续。',
        style='List Bullet',
    )

    doc.add_heading('3.3 箱线图可视化', level=2)
    for part in ['截割部', '牵引部', '油泵', '破碎机']:
        key = f'seg_{part}'
        title, fname = W2_CHART_MAP[key]
        doc.add_heading(title, level=3)
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)

    # ═══ 3.4 异常段检测 ═══
    doc.add_heading('3.4 重采样异常段检测', level=2)
    doc.add_paragraph(
        '重要说明：RLE 段持续时间统计的原始目的不仅是描述操作节奏，'
        '更重要的是检测由 on-change 存储 + 前向填充重采样导致的伪持续段。'
    )
    doc.add_paragraph(
        '数据存储机制：采煤机传感器采用 on-change 存储策略——仅在参数值变化时记录，'
        '不变化时不产生新记录。数据分析时通过前向填充（forward-fill）重采样到 1 分钟'
        '等间隔网格。这意味着：当传感器损坏、断电或通讯中断时，最后上报的值会被'
        '前向填充，形成一段平坦的伪持续段，掩盖了实际的数据缺失。'
    )
    doc.add_paragraph(
        '检测方法：对每个工况状态分别计算段持续时间的分布阈值'
        '（取 max(120min, p99 分位值)），超过阈值且持续 ≥ 60 分钟的段'
        '标记为异常，并根据持续时间长度分级判定原因：'
    )
    anomaly_cause_bullets = [
        '传感器断电/断线：段持续 ≥ 360 分钟（6 小时）',
        '通讯中断/传感器休眠：段持续 ≥ 240 分钟（4 小时）',
        '异常持续段：超过自适应阈值但低于 4 小时',
    ]
    for b in anomaly_cause_bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph(
        '图表解读方法（双面板）：\n'
        '上子图：工况时间线全景图。纵轴为工况类别，横轴为时间。'
        '红色高亮区域为被标记为"异常段"的时段——这些时段内的工况值来自前向填充而非真实传感器上报，'
        '属于伪持续段。\n\n'
        '下子图：段持续时间散点图。每个点代表一个工况连续段，纵轴为该段的持续时间（对数刻度）。'
        '红色点 = 异常段，绿色点 = 正常段。红色虚线为自适应异常阈值'
        '（max(120min, p99 分位值)）。异常段的判定标准：持续超过阈值且时长 ≥ 60 分钟的段，'
        '按分级逻辑判定异常原因。\n\n'
        '右上角标注：异常段数与占比统计，如"15/6654 段，共 4664min = 11.9%"。'
        '占比高（>10%）意味着该部位有大量伪持续数据，阶段二建模时需要特别注意——'
        '这些段内部的真实数据缺失会污染统计基线。'
    )

    # Insert anomaly detection charts
    for part_key, part_label in [('截割部', '截割部'), ('牵引部', '牵引部'),
                                  ('油泵', '油泵'), ('破碎机', '破碎机')]:
        key = f'anomaly_{part_key}'
        title, fname = W2_CHART_MAP.get(key, ('', f'anomalous_segments_{part_key}_工况.png'))
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_heading(f'{part_label}异常段检测（持续时间法）', level=3)
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)

    doc.add_heading('3.5 值异常检测', level=2)
    doc.add_paragraph(
        '除伪持续段检测外，还基于参数数值本身检测异常点/段。'
        '方法：对每个工况状态，以均值 ± 3σ 为阈值标记高偏差点（瞬时异常），'
        '持续 ≥ 10 分钟的连续异常段标记为值异常。'
        '值异常反映传感器真实异常（非数据缺失伪影），是阶段二建模的直接训练目标。'
    )

    for part_key, part_label in [('截割部', '截割部'), ('牵引部', '牵引部'),
                                  ('油泵', '油泵'), ('破碎机', '破碎机')]:
        key = f'value_anomaly_{part_key}'
        title, fname = W2_CHART_MAP.get(key, ('', f'value_anomalies_{part_key}.png'))
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_heading(f'{part_label}值异常检测', level=3)
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 4. 分部位聚类验证规则工况（含规则改进建议）
    # ═══════════════════════════════════════════════════════════
    doc.add_heading('4. 分部位聚类验证规则工况', level=1)

    doc.add_heading('4.1 方法原理', level=2)
    doc.add_paragraph(
        '目的：验证规则工况划分与数据自然聚类的一致性。'
        '如果规则标签（如割煤中/调架中）与无监督聚类结果高度一致，'
        '说明规则阈值的选取符合数据内在结构。'
    )
    doc.add_paragraph(
        '流程：\n'
        '① 选取该部位的关键监测参数作为特征（截割部→截割电流/滚筒高度/摇臂角度；'
        '牵引部→牵引电流/速度；油泵→油泵电流；破碎机→破碎机电流）\n'
        '② StandardScaler 标准化（去均值为 0，方差为 1）\n'
        '③ KMeans 聚类，k = 该部位规则标签类别数\n'
        '④ Adjusted Rand Index（ARI）量化聚类标签与规则标签的一致程度'
    )
    doc.add_paragraph(
        'ARI 取值范围 [-1, 1]：\n'
        '• 1.0 = 完全一致（聚类找到了与规则相同的划分）\n'
        '• 0.0 = 随机分配（聚类结果与规则无关）\n'
        '• 负值 = 比随机还差（极少见，通常表示特征选择有问题）'
    )

    doc.add_heading('4.2 分部位 ARI 结果（优化前基线）', level=2)
    doc.add_paragraph(
        '以下 ARI 基于优化前的工况划分（二值/三态规则）计算。'
        '优化后的多级负载状态（参见 4.4 节）已应用于阶段一最终输出，'
        '但 ARI 定量评估需配合阶段二的重新聚类。'
    )

    ari_data = [
        ['截割部', '0.59', '中',
         '经高度拆分（3m/5m 阈值将割煤中分低位/中位/高位，聚类移除高度特征后 K=7），'
         'ARI 从 0.17 提升至 0.59。KMeans 与规则在高度维度的冲突已消除'],
        ['牵引部', '0.56', '中',
         '空载牵引/重载牵引/待机/停机四态分离较好。电流和速度的边界清晰，'
         '聚类-规则一致性中上'],
        ['油泵', '0.80', '高',
         '停机/轻载/重载三态与自然聚类高度一致——ARI=0.80 说明规则较好地'
         '捕捉了油泵的开关和负载变化模式'],
        ['破碎机', '0.35', '低',
         '三态划分（空载运行/带载运行/停机）虽比二态有提升，但 ARI 仍偏低。'
         '空载运行段中存在多种电流子模式未被规则完全捕捉'],
    ]
    add_table(doc,
        ['部位', 'ARI', '评价', '解释'],
        ari_data,
        col_widths=[2, 1.5, 2, 7],
    )
    spacer(doc)

    doc.add_heading('4.3 可视化', level=2)
    doc.add_paragraph('每个部位一张综合图，左：归一化混淆矩阵，右：PCA 降维散点图：')
    for part in ['截割部', '牵引部', '油泵', '破碎机']:
        key = f'cluster_{part}'
        title, fname = W2_CHART_MAP[key]
        doc.add_heading(title, level=3)
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
        ari_map = {'截割部': '0.59', '牵引部': '0.56', '油泵': '0.80', '破碎机': '0.35'}
        doc.add_paragraph(f'{part}：ARI = {ari_map[part]}。')

    doc.add_heading('4.3.1 特征重要性分析', level=3)
    doc.add_paragraph(
        '为评估每个监测参数对聚类结果的贡献度，使用 Random Forest 分类器'
        '（n_estimators=100, class_weight="balanced"）以聚类标签为目标进行特征重要性排序。'
        '重要性越高表明该参数对区分工况的贡献越大，'
        '反之重要性低的参数在阶段二建模时可考虑移除以降低特征噪声。'
    )
    fi_data = [
        ('截割部', ['左电机温度（37.8%）', '右电机温度（30.5%）', '左摇臂角度（22.8%）', '右摇臂角度（8.9%）']),
        ('牵引部', ['左电机温度（32.8%）', '右电机温度（26.8%）', '采煤机速度（18.5%）', '右电机电流（8.4%）']),
        ('油泵',   ['左电机温度（28.4%）', '右电机温度（24.3%）', '左电机电流（21.5%）', '右电机电流（15.9%）']),
        ('破碎机', ['电机温度（55.9%）', '电机电流（44.1%）']),
    ]
    for part, top_feats in fi_data:
        key = f'fi_{part}'
        title, fname = W2_CHART_MAP[key]
        doc.add_heading(title, level=4)
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 1)
        doc.add_paragraph(
            f'{part}：最重要的特征为 {", ".join(top_feats)}。'
            '重要性排序可为阶段二特征选择提供依据——'
            '排名靠后的特征在异常检测中可直接排除。'
        )

    doc.add_heading('4.3.2 聚类质量评估与优化建议', level=3)
    doc.add_paragraph(
        '综合 ARI（聚类-规则一致性）和 Silhouette Score（簇内紧密度）'
        '对每个部位的聚类配置进行评估，给出阶段二优化建议：'
    )
    quality_data = [
        ['截割部', '0.59', '0.38', '中',
         '经高度拆分（割煤中按 3m/5m 分低位/中位/高位）并从聚类特征中移除高度，'
         'ARI 从 0.17 提升至 0.59。温度（37.8%+30.5%）和摇臂角度（22.8%+8.9%）'
         '成为主导分离特征，规则与数据自然聚类的一致性显著改善。'],
        ['牵引部', '0.56', '0.53', '中',
         '四态工况（空载牵引/重载牵引/待机/停机）聚类一致性中上。'
         '特征重要性显示电机温度（32.8%）和采煤机速度（18.5%）是最强判别指标——'
         '电流的区分力反而低于温度。建议保持现有规则。'],
        ['油泵', '0.80', '0.63', '高',
         '三态工况（轻载/重载/停机）与自然聚类高度一致——'
         'ARI=0.80 说明规则较好地捕捉了油泵的开关和负载变化特征，'
         '轮廓系数 0.63 表示簇内紧密度良好。建议保持现有规则直接用于阶段二建模。'],
        ['破碎机', '0.35', '0.93', '低',
         'ARI 偏低但 Silhouette=0.93 说明自然聚类非常紧密——'
         '空载运行/带载运行/停机三态未能完全匹配聚类分离出的多电流子簇。'
         '破碎机特征维度仅 2 个（电机温度 55.9%, 电机电流 44.1%），'
         '低 ARI 的主因是特征空间过窄而非规则阈值错误。'],
    ]
    add_table(doc,
        ['部位', 'ARI', 'Silhouette', '评价', '优化建议'],
        quality_data,
        col_widths=[1.5, 1, 1, 1.5, 7.5],
    )
    spacer(doc)

    doc.add_heading('4.4 规则优化实施', level=2)
    doc.add_paragraph(
        '基于聚类验证结果，在阶段一收尾阶段对 4 个部位的工况划分规则进行了细化优化。'
        '以下为各部位优化方案（ARI 为优化前参考值）：'
    )
    improvements = [
        ('破碎机工况细化（ARI = 0.35）',
         '在破碎机运行态内增加电机电流阈值 50A，拆分为空载运行（< 50A）和带载运行（≥ 50A）。'
         'Silhouette=0.93 说明自然聚类极紧密但 ARI 仅 0.35——特征空间只有温度和电流 2 维，'
         '工况细分后聚类-规则一致性有提升空间。'),
        ('油泵工况细化（ARI = 0.80）',
         '在油泵运行态内增加平均油压阈值 1.0 MPa，拆分为轻载（< 1.0 MPa）和重载（≥ 1.0 MPa）。'
         'ARI=0.80 说明三态划分高度有效——油泵电流在运行态内存在明显的轻载/重载子模式，'
         '油压比电流更直接反映负载状态。'),
        ('截割部七态划分（ARI = 0.17 → 0.59）',
         '保持速度阈值 0.5 m/min 划分割煤/调架，叠加滚筒高度阈值（3m/5m）将"割煤中"'
         '拆分为低位/中位/高位三子态。同时从聚类特征中移除高度，使规则维度与聚类特征对齐。'
         'ARI 从 0.17 提升至 0.59——KMeans 与规则在高度维度的根本冲突已消除。'
         '温度（37.8%+30.5%）和摇臂角度（22.8%+8.9%）成为主导分离特征。'),
        ('牵引部电流细分（ARI = 0.56）',
         '在"牵引"态内增加左右电机电流之和阈值 < 100A = 空载牵引，≥ 100A = 重载牵引。'
         'ARI=0.56 为中上水平，电机温度（32.8%）是最强区分特征而非电流——'
         '温度的热惯性提供了更稳定的工况基线特征。'),
        ('牵引部三态划分（ARI 最高——基线保留）',
         'ARI 验证其划分最准确。建议保持现行规则，仅在阶段二异常检测中分工况建立基线。'),
    ]
    for title, desc in improvements:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        doc.add_paragraph(desc, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 5. 工况转换时域特征分析与聚合剖面
    # ═══════════════════════════════════════════════════════════
    doc.add_heading('5. 工况转换时域特征分析与聚合剖面', level=1)

    doc.add_heading('5.1 方法原理', level=2)
    doc.add_paragraph(
        '工况转换分析检测工况列的切换事件，提取切换窗口内的时域特征，'
        '量化参数在工况切换时如何变化；同时通过对齐同类切换事件绘制聚'
        '合剖面，揭示参数在切换中的典型变化模式。'
    )

    doc.add_heading('5.1.1 切换检测与特征定义', level=3)
    doc.add_paragraph(
        '基于 shift 比较的切换点检测：\n'
        '• 在工况列上计算 cond.shift(1) ≠ cond → cond ≠ cond.shift(-1)\n'
        '• 定位切换发生的帧\n'
        '• 记录 切换前工况 / 切换后工况 / 切换时刻\n'
        '• 参数提取：使用向量化滚动窗口（.rolling(window=10).mean()）预计算所有窗口均值——'
        'O(1) 复杂度，避免逐行切片。'
    )

    add_table(doc,
        ['特征', '定义', '物理含义'],
        [
            ['Δ_mean', '切换后均值 − 切换前均值',
             '工况切换导致的参数水平变化量（绝对值越大越剧烈）'],
            ['Δ_std_ratio', '|Δ_mean| / min(σₙᵣₑ, σₙₒₛᵥ)',
             '效应量——变化幅度相对于噪声水平的比值（>1 为显著变化）'],
            ['max_slope', '窗口内最大一阶差分',
             '参数变化的最大瞬时速率（反映切换的猛烈程度）'],
            ['rise_time', '从切换点到首次过半程的时间（min）',
             '参数响应速度——电流瞬变 0min，温度滞后数分钟'],
            ['settling_time', '进入 [后均值±5%] 后不再离开的时间',
             '达到新稳态所需时间（若在窗口内未稳定则 = 窗口长度）'],
            ['overshoot', '(峰值 − 后均值) / Δ_mean',
             '超调量——参数越过新稳态的程度（>0.1 = 存在明显超调）'],
            ['energy_ratio', '后窗 RMS / 前窗 RMS',
             '切换前后信号能量比（电流变换时能量可翻倍）'],
        ],
        col_widths=[2.5, 4.5, 5.5],
    )
    spacer(doc)

    doc.add_heading('5.1.2 聚合剖面方法', level=3)
    doc.add_paragraph(
        '聚合剖面（Aggregate Transition Profile）：将同类型的所有切换事件按切换时刻'
        '对齐（t=0），叠加各自的参数轨迹（灰色细线），绘制均值±1σ 包络（橙色阴影+暗红线）。'
        '一次聚合显示一个参数在一种切换类型下的典型变化模式。'
    )
    doc.add_paragraph(
        '通过聚合剖面可以直观回答：电流在割煤中→正常运行切换时走什么形状？'
        '是阶跃、斜坡还是过冲？每次是否一致？'
    )

    doc.add_heading('5.2 设备级切换类型分布', level=2)
    doc.add_paragraph('共检测到 6,813 次设备级工况切换事件，分布如下：')

    add_table(doc,
        ['切换类型', '频次', '占比'],
        [
            ['割煤中 ↔ 调架中', '4,453', '65.4%'],
            ['割煤中 ↔ 待机',   '1,228', '18.0%'],
            ['调架中 ↔ 待机',     '689', '10.1%'],
            ['待机 ↔ 停机',       '238',  '3.5%'],
            ['空载牵引相关',       '99',  '1.5%'],
            ['其他',              '106',  '1.6%'],
        ],
        col_widths=[4, 2.5, 2],
    )
    spacer(doc)
    doc.add_paragraph(
        '割煤中↔调架中互切占 65.4%，这是采煤机单刀作业的典型循环'
        '——每割完一刀在工作面两端切换方向，期间需要调架换向。',
        style='List Bullet',
    )

    doc.add_heading('5.3 关键时域特征发现', level=2)

    doc.add_heading('上升时间（rise_time）—— 热惯性量化', level=3)
    add_table(doc,
        ['参数', 'p50（min）', 'p75（min）', 'p90（min）', '响应特征'],
        [
            ['采煤机速度',   '0', '1', '2', '机械响应：几乎瞬变'],
            ['截割电流（右）', '0', '1', '3', '电气响应：1 分钟窗口内完成阶跃'],
            ['截割电流（左）', '0', '1', '3', '电气响应：与右滚筒对称'],
            ['截割温度（右）', '0', '2', '5', '热惯性：p90 延迟达 5 分钟'],
            ['截割温度（左）', '0', '2', '4', '热惯性：p90 延迟达 4 分钟'],
            ['牵引电流（右）', '0', '1', '2', '电气响应：快速阶跃'],
            ['牵引温度（右）', '0', '2', '4', '热惯性：p90 延迟达 4 分钟'],
        ],
        col_widths=[3, 2, 2, 2, 4.5],
    )
    spacer(doc)
    doc.add_paragraph(
        '核心发现：电流/速度的上升时间 p50 = 0 分钟（瞬态响应），'
        '而温度的 p90 = 4~5 分钟——热惯性导致温度响应滞后电流变化 3~5 分钟。'
        '这种滞后在处理电流-温度联合特征时必须考虑时间偏移补偿。',
        style='List Bullet',
    )

    doc.add_heading('效应量（Δ_std_ratio）—— 参数对工况切换的敏感度', level=3)
    add_table(doc,
        ['参数类型', '|Δ|/std > 1 占比', '中位数效应量', '敏感度评级'],
        [
            ['速度',       '21%', '0.46', '低 —— 很多切换前后速度不变'],
            ['电流',       '27~33%', '0.51~0.62', '中 —— 部分切换电流变化显著'],
            ['温度',       '53~69%', '1.14~1.82', '高 —— 温度是最佳工况区分器'],
        ],
        col_widths=[3, 3.5, 3, 4],
    )
    spacer(doc)
    doc.add_paragraph(
        '温度是比电流更稳定的工况区分器：电流在割煤中内部的散布大，'
        '很多切换的电流效应量 < 1；温度变化虽慢但幅度大、信噪比高。'
        '推荐在异常检测中将温度作为工况稳定性监测的关键指标。',
        style='List Bullet',
    )

    doc.add_heading('超调量（overshoot）—— 切换冲击程度', level=3)
    doc.add_paragraph(
        '92~94% 的电流/速度切换事件存在明显超调（overshoot > 0.1），'
        '均值超调量达 9~15 倍。这意味着每次工况切换时，电流会越过新稳态值'
        '再回落——这对电机和电气系统构成反复的电流冲击。'
    )
    doc.add_paragraph(
        '温度的超调占比 80~94%，但均值仅 2~3 倍——温度有热惯性自然平滑，'
        '超调幅度远小于电流。',
        style='List Bullet',
    )

    doc.add_heading('5.4 聚合剖面可视化', level=2)
    for label_key in ['profile_电流', 'profile_速度', 'profile_温度']:
        title, fname = W2_CHART_MAP[label_key]
        doc.add_heading(title, level=3)
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
        descs = {
            'profile_电流': (
                '电流聚合剖面：割煤中→调架中（上）电流从工作值骤降至接近 0，'
                '伴随大幅超调（灰色细线在 t=0 后散布极宽）；'
                '待机→割煤中（下）电流从 0 阶跃至工作值，上升沿陡峭。'
            ),
            'profile_速度': (
                '速度聚合剖面：割煤中→调架中速度从 2~6 m/min 降至接近 0，'
                '下降沿斜率一致性好（灰色细线聚集）；'
                '割煤中→待机速度直接归零。'
            ),
            'profile_温度': (
                '温度聚合剖面：与电流的瞬态响应不同，温度在 t=0 后缓慢上升，'
                '5~10 分钟才达到新稳态。灰色细线的散布较小，说明温度变化模式'
                '比电流更一致（热容量的低通滤波效应）。'
            ),
        }
        doc.add_paragraph(descs.get(label_key, ''))

    doc.add_heading('5.5 多参数切换签名（热力图指纹）', level=2)
    doc.add_paragraph(
        '多参数切换签名：以热力图形式展示一次切换事件中多个参数的 z-score 协同变化。'
        'X 轴 = 切换窗口相对时间 [-15min, +15min]，Y 轴 = 参数按影响大小排序，'
        '色阶 = z-score（红升蓝降），右侧柱状图 = Δ_mean。'
        '对每种参数组合（电流+速度、电流+温度、全关键参数），取 top-3 最频繁切换类型的'
        '首个事件生成指纹图，展示不同切换模式下参数的差异化响应。'
    )

    for label_key in sorted([k for k in W2_CHART_MAP if k.startswith('multi_')]):
        title, fname = W2_CHART_MAP[label_key]
        doc.add_heading(title, level=3)
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
            # Extract description from key
            suffix_part = label_key.replace('multi_', '').rsplit('_', 1)[0]
            type_idx = label_key.rsplit('_', 1)[-1]
            doc.add_paragraph(
                f'第 {int(type_idx)+1} 大切换类型的首事件指纹——'
                f'参数组：{suffix_part}，热力图中红色=参数升高，蓝色=降低，'
                f'竖向红色虚线=t=0 切换时刻。'
            )
        else:
            doc.add_paragraph(f'（无图表数据: {fname}）')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 6. 分部位滞后互相关分析
    # ═══════════════════════════════════════════════════════════
    doc.add_heading('6. 分部位滞后互相关分析', level=1)
    doc.add_paragraph(
        '前文设备级滞后互相关（电流 vs 温度 / 电流 vs 速度 / 俯仰角 vs 位置架号）'
        '关注的是整机宏观耦合，但不同部位的结构和监测参数差异显著，'
        '设备级结果可能掩盖部位特有的动态关系。'
    )
    doc.add_paragraph(
        '因此进一步按部位细分：每个部位选出最相关的两个参数对，'
        '分别计算 -60 ~ +60 分钟的滞后互相关，定位最优滞后时间（最佳延迟）和相关系数。'
    )

    # ── 6.1 截割部 ──
    doc.add_heading('6.1 截割部：电流 vs 温度', level=2)
    if 'lag_cmj_截割部' in W2_CHART_MAP:
        title, fname = W2_CHART_MAP['lag_cmj_截割部']
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
    doc.add_paragraph(
        '截割部电流与温度的滞后互相关揭示：电流变化在前，温度响应滞后约 10~20 分钟。'
        '这是截割部传动系统热惯性的直接量化——割煤时电流增大产生热量，'
        '温度在数分钟后才达到峰值。该滞后可作为阶段二温度异常检测的时间偏移补偿依据。'
    )

    # ── 6.2 牵引部 ──
    doc.add_heading('6.2 牵引部：电流 vs 速度', level=2)
    if 'lag_cmj_牵引部' in W2_CHART_MAP:
        title, fname = W2_CHART_MAP['lag_cmj_牵引部']
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
    doc.add_paragraph(
        '牵引部电流与速度的耦合最紧密（电机的电磁-机械响应在秒级完成），'
        '最优滞后通常 ≤ 5 分钟。若滞后超过 10 分钟，可能反映牵引电机励磁响应延迟或传动链间隙。'
    )

    # ── 6.3 油泵 ──
    doc.add_heading('6.3 油泵：电流 vs 油压', level=2)
    if 'lag_cmj_油泵' in W2_CHART_MAP:
        title, fname = W2_CHART_MAP['lag_cmj_油泵']
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
    doc.add_paragraph(
        '油泵电流与输出油压之间理论上呈正相关——电流增大驱动油压升高。'
        '实际滞后互相关可指示液压系统的响应速度：滞后过长可能反映液压阀动作延迟或油路泄漏。'
    )

    # ── 6.4 破碎机 ──
    doc.add_heading('6.4 破碎机：电流 vs 温度', level=2)
    if 'lag_cmj_破碎机' in W2_CHART_MAP:
        title, fname = W2_CHART_MAP['lag_cmj_破碎机']
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
    doc.add_paragraph(
        '破碎机电流与温度同样存在热惯性滞后，但与截割部相比，破碎机负载变化更剧烈'
        '（煤块瞬时冲击），电流-温度相关性可能较低（破碎作功不完全转化为热能）。'
    )

    # ── 6.5 转载机 ──
    doc.add_heading('6.5 转载机：电流 vs 转速', level=2)
    if 'lag_zzj' in W2_CHART_MAP:
        title, fname = W2_CHART_MAP['lag_zzj']
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
    doc.add_paragraph(
        '转载机电流-转速滞后互相关反映传动链的机电耦合——'
        '变频器调节转速后电流随之变化。'
        '转载机采用变频驱动，其电流-转速响应通常比工频驱动的采煤机更快。'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 7. 转载机深度分析
    # ═══════════════════════════════════════════════════════════
    doc.add_heading('7. 转载机深度分析', level=1)
    doc.add_paragraph(
        '转载机（ZZJ）是工作面运输系统的核心设备，负责将采煤机割落的原煤'
        '转运至顺槽皮带。与采煤机不同，转载机不涉及截割动作，'
        '其工况仅通过一条 "工况" 列进行三态划分（停机 / 空载运行 / 带载运行）。'
        '以下对转载机进行与采煤机对等的深度分析。'
    )

    # ── 7.1 Kruskal-Wallis ──
    doc.add_heading('7.1 Kruskal-Wallis 非参数检验', level=2)
    doc.add_paragraph(
        '对转载机各监测参数在三种工况（停机/空载运行/带载运行）之间的分布差异性'
        '进行 Kruskal-Wallis 检验，使用 FDR 校正 p 值，'
        '以 ε² 作为效应量评估差异的实际显著程度。'
    )
    if 'kw_转载机' in W2_CHART_MAP:
        title, fname = W2_CHART_MAP['kw_转载机']
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_heading(title, level=3)
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
    doc.add_paragraph(
        '转载机 KW 检验可验证 "空载/带载运行" 的划分是否在各参数维度上均有显著区分。'
        '若某参数在空载和带载间无显著差异（p > 0.05 且 ε² < 0.01），'
        '可能意味着该参数对负载不敏感，不适合作为后续异常检测的特征。'
    )

    # ── 7.2 段持续与异常段 ──
    doc.add_heading('7.2 段持续时间与异常段检测', level=2)
    doc.add_paragraph(
        '转载机同样基于 on-change 存储机制，重采样后存在伪持续段问题。'
        '采用与采煤机相同的 RLE 段持续时间分析方法，'
        '检测转载机工况的异常持续段和值异常。'
    )
    for seg_key in ['seg_转载机', 'anomaly_转载机']:
        if seg_key in W2_CHART_MAP:
            title, fname = W2_CHART_MAP[seg_key]
            fpath = os.path.join(PHASE1, fname)
            if os.path.exists(fpath):
                doc.add_heading(title, level=3)
                doc.add_picture(fpath, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                spacer(doc, 2)
    doc.add_paragraph(
        '转载机工况段持续时间分析揭示：转载机长期处于带载运行状态（占比约 59.2%），'
        '停机段较短且分散；空载运行段占比极低（约 1.1%）。异常段检测可识别因通讯中断'
        '或传感器休眠导致的伪持续段，辅助评估数据质量。'
    )

    # ── 7.3 聚类验证 ──
    doc.add_heading('7.3 聚类验证规则工况', level=2)
    doc.add_paragraph(
        '以转载机关键监测参数（电流、转速、转矩等）为特征，'
        '使用 KMeans 进行无监督聚类，将聚类结果与规则工况标签对比，'
        '通过 ARI（调整兰德指数）和轮廓系数评估规则划分与数据自然聚类的一致性。'
    )
    for cl_key in ['cluster_转载机', 'fi_转载机']:
        if cl_key in W2_CHART_MAP:
            title, fname = W2_CHART_MAP[cl_key]
            fpath = os.path.join(PHASE1, fname)
            if os.path.exists(fpath):
                doc.add_heading(title, level=3)
                doc.add_picture(fpath, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                spacer(doc, 2)
    doc.add_paragraph(
        '转载机规则工况（停机/空载运行/带载运行）的三态划分较为明确，'
        'ARI 通常较高（> 0.8），说明规则阈值选取合理。'
        '特征重要性分析可识别对工况分类贡献最大的监测参数，'
        '为阶段二异常检测的特征筛选提供依据。'
    )

    # ── 7.4 工况转换 ──
    doc.add_heading('7.4 工况转换分析', level=2)
    doc.add_paragraph(
        '检测转载机工况切换事件（停机→空载、空载→带载等），'
        '统计切换频率、方向分布，并分析切换前后各监测参数的均值变化（Δ），'
        '评估各参数对工况切换的敏感度。'
    )
    if 'transition_zzj' in W2_CHART_MAP:
        title, fname = W2_CHART_MAP['transition_zzj']
        fpath = os.path.join(PHASE1, fname)
        if os.path.exists(fpath):
            doc.add_heading(title, level=3)
            doc.add_picture(fpath, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            spacer(doc, 2)
    doc.add_paragraph(
        '转载机工况转换频率显著低于采煤机各部位——转载机一旦进入带载运行，'
        '通常持续较长时间（数小时），不会像截割部那样频繁切换。'
        '低切换频率意味着转载机数据中的暂态段较少，稳态段占主导，'
        '有利于建立稳定的工况基线。'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # 8. 关键发现与阶段二预备
    # ═══════════════════════════════════════════════════════════
    doc.add_heading('8. 关键发现与阶段二预备', level=1)

    doc.add_heading('8.1 综合分析结论', level=2)

    findings = [
        ('数据质量极高',
         '全体 21 个监测参数覆盖率 ≥ 99.86%，无需数据质量层面的预处理和插补。'
         '这是后续建模的坚实基础。'),
        ('分部位工况划分有效：Kruskal-Wallis 全部显著',
         '所有参数在不同工况间分布差异显著（FDR 校正 p < 0.05）。'
         '速度、电流是最强的工况区分器，H 统计量最高达 33,350（牵引部速度）；'
         '俯仰角几乎不受工况影响（H 统计量最低）。'),
        ('异常段检测揭示数据缺失模式',
         'on-change 存储 + 前向填充重采样在前端状态下产生了伪持续段。'
         '异常段检测为阶段二的数据清洗提供了精确的目标——在建模前截断或排除这些伪段。'),
        ('调架中短促频繁，割煤中段持续中位 4 分钟——操作节奏反映工艺',
         '调架中段（截割部）持续中位仅 1 分钟，段数 2,588 段，说明调架短促频繁。'
         '割煤中段中位 4 分钟，最大 414 分钟≈6.9 小时。'),
        ('规则优化已实施：分部位工况从二值/三态细化为多级负载状态',
         '基于 ARI 分析和 KW 显著参数，对 4 个部位实施了工况优化——'
         '截割部：割煤中/调架中/待机-高位/待机/停机（ARI=0.17→预期提升，'
         '通过左滚筒高度阈值 4.5m 拆分待机态，使规则与聚类特征对齐）；'
         '牵引部：空载/重载牵引（ARI=0.56）；'
         '油泵：轻载/重载（油压阈值 1.0 MPa，ARI=0.80）；'
         '破碎机：空载/带载运行（电流阈值 50A，ARI=0.35）。'
         '新规则在阶段二异常检测中将提供更精细的工况基线。'),
        ('电流-温度异步响应：热惯性产生 3~5 分钟滞后',
         '电流 p90 上升时间 2~3 分钟，温度 p90 上升时间 4~5 分钟。'
         '跨参数分析要考虑时间偏移补偿。'),
        ('电机温度是最强聚类区分特征',
         'Random Forest 特征重要性显示：牵引部左电机温度（32.8%）、'
         '油泵左电机温度（28.4%）、破碎机电机温度（55.9%）均为各自部位的第一区分特征。'
         '建议将温度纳入阶段二异常检测的核心特征集。'),
        ('92%+ 切换事件存在电流超调',
         '每次工况切换都对电机和电网形成电流冲击。'
         '超调幅度可作为电气系统健康状态的间接指标。'),
        ('转载机深度分析完成：KW 显著、ARI 高、切换频率低',
         '转载机所有参数在三种工况间均有显著差异（KW p < 0.05）；'
         '聚类验证 ARI=0.85 表明规则工况与数据自然聚类高度一致；'
         '工况切换频率远低于采煤机——稳态段占主导，有利于建立稳定的工况基线。'
         '转载机同样存在 on-change 存储导致的伪持续段（0.8% 段 + 18.1% 时长），'
         '阶段二建模时需注意截断或排除这些伪段。'),
    ]
    for title, desc in findings:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}')
        run.font.bold = True
        doc.add_paragraph(desc, style='List Bullet')

    doc.add_heading('8.2 阶段二异常检测设计建议', level=2)

    doc.add_paragraph('基于第二周分析结果，阶段二的异常检测策略应包含以下要素：')

    add_table(doc,
        ['维度', '指标设计', '依据'],
        [
            ['分工况基线',
             '每个参数在每种工况下的 mean±3σ / 中位±IQR',
             'Kruskal-Wallis 证实工况间的参数分布差异显著，'
             '必须分工况计算基线'],
            ['温升率监控',
             '温度一阶导数的滑动窗口标准差',
             '温度变化慢但信噪比高，温升率异常可预警散热故障'],
            ['电流-温度异步检测',
             '电流上升后 N 分钟内温度是否跟进',
             '热惯性产生 3~5 分钟滞后，延迟偏离正常范围→传感器/散热异常'],
            ['切换冲击累计',
             '单位时间内 overshoot > 阈值的切换次数',
             '92% 切换有超调，高频冲击累计可能导致电气系统疲劳'],
            ['异常段截断处理',
             '在建模前根据异常段检测结果截断或排除 6h+ 的伪持续段',
             'on-change 前向填充会掩盖真实的数据缺失，'
             '异常段在阶段二作为缺失值处理，避免建模偏差'],
            ['破碎机工况细化',
             '在运行态内做 KMeans 子聚类 → 轻/中/重载分离',
             'ARI = 0.10 说明二值规则太粗，聚类能发现更细的异常子模式'],
        ],
        col_widths=[3, 4.5, 5],
    )
    spacer(doc)

    doc.add_paragraph(
        '阶段二异常检测方法的实施优先级：异常段截断清洗 → 3σ/IQR 基线 '
        '→ 滑动窗口温升率监控 → 电流-温度异步检测 → Isolation Forest 补充。'
        '优先级设计的核心逻辑：物理可解释的方法优先，黑箱方法作为补充。'
    )

    # ── 保存 ──
    os.makedirs(PHASE1, exist_ok=True)
    doc.save(W2_OUTPUT_FILE)
    print(f'[OK] 第二周报告已生成: {W2_OUTPUT_FILE}')
    print(f'    文件大小: {os.path.getsize(W2_OUTPUT_FILE) / 1024:.0f} KB')



if __name__ == '__main__':
    build_report()
